#!/usr/bin/env python3
"""
Dopunjuje rad poglavljima koja nedostaju za metode dodane nakon prve verzije.

Pokrece se NAKON azuriraj_word.py i radi nad njegovim izlazom. Dodaje:
  - teorijsko potpoglavlje o neuronskoj mrezi,
  - implementacijsko potpoglavlje o neuronskoj mrezi,
  - implementacijsko potpoglavlje o zajednickim znacajkama ML metoda,
te uskladuje opise naprednog KNN-a, pomicnog prosjeka, stabla odlucivanja i
slucajne sume s onim sto kod stvarno radi. Dopunjuje i literaturu, popis
kratica, sazetak i kljucne rijeci.

Novi odlomci nastaju kloniranjem postojecih, pa preuzimaju njihov stil u
cijelosti. Skripta je idempotentna: ako potpoglavlja vec postoje, preskace ih.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, str(Path(__file__).resolve().parent))
from azuriraj_word import norm, podesi_broj_redaka, upisi_celiju, upisi_odlomak  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
DOK = ROOT / "docs" / "Diplomski-Toni_Jakelic_azurirano.docx"

promasaji: list[str] = []
brojac = {"zamjena": 0, "umetnuto": 0}


# --------------------------------------------------------------------------
def nadi(doc, tekst: str):
    cilj = norm(tekst)
    for par in doc.paragraphs:
        if norm(par.text) == cilj:
            return par
    return None


def nadi_pocetak(doc, prefiks: str):
    cilj = norm(prefiks)
    for par in doc.paragraphs:
        if norm(par.text).startswith(cilj):
            return par
    return None


def zamijeni(doc, stari: str, novi: str, oznaka: str) -> None:
    par = nadi(doc, stari)
    if par is None:
        promasaji.append(oznaka)
        return
    upisi_odlomak(par, novi)
    brojac["zamjena"] += 1


def klon_iza(uzor, sidro, tekst: str):
    """
    Umece kopiju odlomka 'uzor' odmah iza 'sidro' i upisuje zadani tekst.

    Iz kopije se uklanjaju oznake i polja jer bi udvostrucena imena knjiznih
    oznaka pokvarila kazalo i poveznice u sadrzaju.
    """
    el = copy.deepcopy(uzor._element)
    for tag in ("w:bookmarkStart", "w:bookmarkEnd", "w:fldSimple",
                "w:fldChar", "w:instrText", "w:proofErr", "w:commentRangeStart",
                "w:commentRangeEnd"):
        for e in list(el.iter(qn(tag))):
            roditelj = e.getparent()
            if roditelj is not None:
                roditelj.remove(e)
    sidro._element.addnext(el)
    novi = Paragraph(el, uzor._parent)
    upisi_odlomak(novi, tekst)
    brojac["umetnuto"] += 1
    return novi


def umetni_odjeljak(doc, sidro, naslov_uzor, tijelo_uzor,
                    naslov: str, odlomci: list[str], oznaka: str) -> None:
    if nadi(doc, naslov) is not None:
        print(f"  preskacem, vec postoji: {naslov}")
        return
    if sidro is None or naslov_uzor is None or tijelo_uzor is None:
        promasaji.append(oznaka)
        return
    tekuci = klon_iza(naslov_uzor, sidro, naslov)
    for tekst in odlomci:
        tekuci = klon_iza(tijelo_uzor, tekuci, tekst)


# --------------------------------------------------------------------------
def main() -> None:
    if not DOK.exists():
        sys.exit(f"Nema {DOK}. Prvo pokreni scripts/azuriraj_word.py")
    doc = docx.Document(str(DOK))

    # Uzorci iz kojih se kloniraju novi odlomci.
    h_impl = nadi(doc, "Implementacija Random Forest metode")       # Heading 2
    h_teor = nadi(doc, "Random Forest metoda")                      # Heading 3
    tijelo = nadi_pocetak(doc, "Cijena te stabilnosti je veća složenost")
    if h_impl is None or h_teor is None or tijelo is None:
        sys.exit("Ne mogu pronaci uzorke za kloniranje.")
    print(f"uzorci: naslov2={h_impl.style.name}, naslov3={h_teor.style.name}, "
          f"tijelo={tijelo.style.name}")

    # ---------------- TEORIJA: uvod u metode strojnog ucenja -----------------
    zamijeni(
        doc,
        "Metode strojnog učenja u ovom radu koriste poznate dijelove temperaturnog niza kako bi "
        "procijenile vrijednosti koje su umjetno uklonjene. Za razliku od jednostavnih "
        "interpolacijskih metoda, one mogu koristiti dodatne značajke, primjerice položaj mjerenja u "
        "nizu, sat u danu ili dan u godini. Opći postupak sastoji se od izdvajanja poznatih uzoraka, "
        "izračuna značajki, učenja odnosa između značajki i temperature te predikcije nedostajućih "
        "vrijednosti.",
        "Metode strojnog učenja u ovom radu koriste poznate dijelove temperaturnog niza kako bi "
        "procijenile vrijednosti koje su umjetno uklonjene. Za razliku od jednostavnih "
        "interpolacijskih metoda, one mogu istodobno koristiti više izvora informacija: vrijednosti na "
        "rubovima praznine i udaljenost do njih, položaj mjerenja u nizu te sat u danu i dan u godini. "
        "Opći postupak sastoji se od izdvajanja poznatih uzoraka, izračuna značajki, učenja odnosa "
        "između značajki i temperature te predikcije nedostajućih vrijednosti. Sve metode strojnog "
        "učenja u ovom radu uče odstupanje od linearne procjene umjesto same temperature, pa im je "
        "linearna interpolacija zajednička polazna točka koju pokušavaju popraviti.",
        "teorija: uvod u ML",
    )

    # ---------------- TEORIJA: pomicni prosjek -------------------------------
    zamijeni(
        doc,
        "U završnom eksperimentalnom CSV-u metoda pomičnog prosjeka nije uključena u usporedbu "
        "rezultata. U teorijskom dijelu opisana je kao tipična klasična metoda, dok su u "
        "implementiranom eksperimentu korištene forward fill, linearna, vremenska, kubna i spline "
        "interpolacija.",
        "U završnom eksperimentu pomični prosjek uspoređen je zajedno s ostalim metodama, uz prozor od "
        "šest uzoraka s obje strane praznine, što pri desetominutnom razmaku odgovara jednom satu. "
        "Njegovo se ponašanje jasno razdvaja po scenarijima. Kod pojedinačno raspoređenih praznina daje "
        "osrednji rezultat jer zaglađivanje uklanja i stvarne promjene signala, a ne samo šum. U "
        "unutrašnjosti duljih blokova unutar prozora nema nijednog poznatog mjerenja, pa se metoda "
        "svodi na zadržavanje posljednje poznate vrijednosti i daje gotovo iste rezultate kao forward "
        "fill.",
        "teorija: pomicni prosjek",
    )

    # ---------------- TEORIJA: napredni KNN ----------------------------------
    zamijeni(
        doc,
        "Napredna KNN varijanta mijenja način na koji su ulazne značajke predstavljene prije izračuna "
        "udaljenosti. Položaj mjerenja u nizu normalizira se kako njegova brojčana veličina ne bi "
        "potisnula ostale informacije, dok se sat u danu i dan u godini prikazuju kao cikličke "
        "značajke. Ciklički prikaz posebno je važan za vrijeme. Primjerice, 23:00 i 00:00 nalaze se "
        "jedan sat jedno od drugoga, iako ih obične numeričke vrijednosti 23 i 0 prikazuju kao vrlo "
        "udaljene. Sinusna i kosinusna komponenta uklanjaju taj umjetni prekid i povezuju kraj ciklusa "
        "s njegovim početkom.",
        "Napredna KNN varijanta mijenja ono što se uspoređuje. Umjesto da traži mjerenja obavljena u "
        "blizini nedostajuće točke, ona traži mjerenja koja su se nalazila u sličnoj situaciji. "
        "Situacija se opisuje relativnim položajem točke unutar praznine, udaljenostima do poznatih "
        "vrijednosti s njezine lijeve i desne strane te satom u danu prikazanim cikličkim komponentama. "
        "Ciklički prikaz posebno je važan za vrijeme. Primjerice, 23:00 i 00:00 nalaze se jedan sat "
        "jedno od drugoga, iako ih obične numeričke vrijednosti 23 i 0 prikazuju kao vrlo udaljene. "
        "Sinusna i kosinusna komponenta uklanjaju taj umjetni prekid i povezuju kraj ciklusa s njegovim "
        "početkom.",
        "teorija: napredni KNN prvi odlomak",
    )

    zamijeni(
        doc,
        "Pri završnoj procjeni susjedi nemaju jednaku težinu. Uzorci koji su prema odabranim značajkama "
        "bliži nedostajućoj točki više utječu na rezultat, dok udaljeniji imaju manji doprinos. Time se "
        "procjena više oslanja na vremenski slične dijelove niza.",
        "Druga bitna razlika je u tome što napredna varijanta ne procjenjuje samu temperaturu, nego "
        "odstupanje od jednostavne linearne procjene. Time polazi od pretpostavke da je linearna "
        "procjena dobra osnova i pokušava naučiti u kojim se situacijama ona sustavno vara i za koliko. "
        "Pri završnom izračunu susjedi nemaju jednaku težinu: uzorci čija je situacija sličnija više "
        "utječu na rezultat, dok udaljeniji imaju manji doprinos.",
        "teorija: napredni KNN drugi odlomak",
    )

    # Tablica 4-1: usporedba osnovnog i naprednog KNN-a.
    t_knn = doc.tables[0]
    redovi_knn = [
        ["Što se uspoređuje", "blizina u samom nizu", "sličnost situacije u kojoj je praznina"],
        ["Ulazne veličine", "položaj mjerenja u nizu",
         "relativan položaj u praznini, udaljenost do oba ruba, sat u danu"],
        ["Odabir susjeda", "po jedan poznati susjed sa svake strane praznine",
         "dvanaest poznatih točaka s najsličnijom situacijom"],
        ["Što se procjenjuje", "sama temperatura", "odstupanje od linearne procjene"],
        ["Težine susjeda", "inverzna udaljenost u nizu",
         "inverzna udaljenost u prostoru opisa situacije"],
        ["Cilj", "pouzdana lokalna procjena",
         "popravak linearne procjene ondje gdje ona griješi"],
    ]
    podesi_broj_redaka(t_knn, len(redovi_knn))
    for i, red in enumerate(redovi_knn, start=1):
        for j, v in enumerate(red):
            upisi_celiju(t_knn.rows[i].cells[j], v)
    print(f"  tablica 4-1 popunjena: {len(redovi_knn)} redaka")

    # ---------------- TEORIJA: novo potpoglavlje o neuronskoj mrezi ----------
    umetni_odjeljak(
        doc,
        sidro=nadi_pocetak(doc, "Cijena te stabilnosti je veća složenost"),
        naslov_uzor=h_teor,
        tijelo_uzor=tijelo,
        naslov="Neuronska mreža",
        odlomci=[
            "Neuronska mreža procjenjuje temperaturu nizom težinskih zbrajanja ulaznih značajki i "
            "nelinearnih preslikavanja. Ulazne vrijednosti prolaze kroz jedan ili više skrivenih "
            "slojeva, u kojima svaki neuron računa težinski zbroj svojih ulaza i na njega primjenjuje "
            "nelinearnu funkciju, a izlazni sloj taj rezultat pretvara u konačnu procjenu. [15]",
            "Za razliku od stabala odlučivanja, koja prostor značajki dijele pravokutnim rezovima, "
            "mreža može naučiti glatke i postupne prijelaze. To načelno odgovara temperaturi, koja se "
            "mijenja kontinuirano. Učenje se provodi postupnim smanjivanjem pogreške: mreža za poznate "
            "uzorke usporedi svoju procjenu sa stvarnom vrijednošću, propagacijom unatrag izračuna "
            "koliko je svaka težina pridonijela toj pogrešci i pomakne težine u smjeru koji pogrešku "
            "smanjuje.",
            "Cijena te fleksibilnosti je potreba za većom količinom podataka i osjetljivost na postavke "
            "učenja, poput broja neurona, stope učenja i broja prolazaka kroz podatke. Mreža također ne "
            "nudi objašnjenje pojedinačne procjene, jer je ona posljedica velikog broja težina koje se "
            "ne mogu pročitati kao skup razumljivih pravila.",
        ],
        oznaka="teorija: neuronska mreza",
    )

    # ---------------- IMPLEMENTACIJA: klasicne metode ------------------------
    zamijeni(
        doc,
        "Klasične metode implementirane su kao skup funkcija koje nad istim oštećenim temperaturnim "
        "nizom stvaraju rekonstruirani izlazni niz. Ulaz u svaku metodu je niz u kojem su pojedine "
        "vrijednosti označene kao NaN, dok je izlaz novi niz u kojem se te vrijednosti zamjenjuju "
        "procjenama. Poznate vrijednosti pritom se zadržavaju, kako bi se mijenjale samo vrijednosti "
        "koje su stvarno nedostajale. Forward fill metoda prolazi kroz niz i svaku nedostajuću "
        "vrijednost popunjava zadnjom prethodno poznatom temperaturom. Linearna interpolacija pronalazi "
        "lijevu i desnu poznatu vrijednost oko praznine te vrijednosti između njih računa kao "
        "ravnomjernu promjenu. Vremenska interpolacija koristi isti princip, ali umjesto samog rednog "
        "broja uzorka uzima u obzir stvarnu vremensku oznaku. Kubna i spline interpolacija koriste "
        "zakrivljene funkcije za procjenu nedostajućih vrijednosti te mogu bolje pratiti nelinearne "
        "promjene u nizu. Njihovo ograničenje je veća osjetljivost na rubne slučajeve i dulje blokove "
        "nedostajućih podataka. U završnom eksperimentu uspoređeno je pet klasičnih metoda: forward "
        "fill, linearna interpolacija, vremenska interpolacija, kubna interpolacija i spline "
        "interpolacija.",
        "Klasične metode implementirane su kao skup funkcija koje nad istim oštećenim temperaturnim "
        "nizom stvaraju rekonstruirani izlazni niz. Ulaz u svaku metodu je niz u kojem su pojedine "
        "vrijednosti označene kao NaN, dok je izlaz novi niz u kojem se te vrijednosti zamjenjuju "
        "procjenama. Poznate vrijednosti pritom se zadržavaju, kako bi se mijenjale samo vrijednosti "
        "koje su stvarno nedostajale. Forward fill metoda prolazi kroz niz i svaku nedostajuću "
        "vrijednost popunjava zadnjom prethodno poznatom temperaturom. Linearna interpolacija pronalazi "
        "lijevu i desnu poznatu vrijednost oko praznine te vrijednosti između njih računa kao "
        "ravnomjernu promjenu. Vremenska interpolacija koristi isti princip, ali umjesto samog rednog "
        "broja uzorka uzima u obzir stvarnu vremensku oznaku. Kubna i spline interpolacija koriste "
        "zakrivljene funkcije za procjenu nedostajućih vrijednosti te mogu bolje pratiti nelinearne "
        "promjene u nizu. Njihovo ograničenje je veća osjetljivost na rubne slučajeve i dulje blokove "
        "nedostajućih podataka. Pomični prosjek za svaku prazninu prikuplja poznata mjerenja unutar "
        "prozora od šest uzoraka lijevo i desno te računa njihov prosjek; kada unutar tog prozora nema "
        "nijednog poznatog mjerenja, što se događa u unutrašnjosti duljih blokova, vrijednost se "
        "popunjava posljednjom poznatom temperaturom. U završnom eksperimentu uspoređeno je šest "
        "klasičnih metoda: forward fill, linearna interpolacija, vremenska interpolacija, kubna "
        "interpolacija, spline interpolacija i pomični prosjek.",
        "implementacija: klasicne metode",
    )

    # ---------------- IMPLEMENTACIJA: zajednicke znacajke --------------------
    umetni_odjeljak(
        doc,
        sidro=nadi_pocetak(doc, "Klasične metode implementirane su kao skup funkcija"),
        naslov_uzor=h_impl,
        tijelo_uzor=tijelo,
        naslov="Značajke za metode strojnog učenja",
        odlomci=[
            "Sve metode strojnog učenja dijele isti skup ulaznih značajki, pa razlika u njihovim "
            "rezultatima proizlazi iz modela, a ne iz različito pripremljenih ulaza. Za svaku točku "
            "niza računa se jedanaest značajki. Njih šest opisuje prazninu u kojoj se točka nalazi: "
            "posljednja poznata vrijednost s lijeve strane, prva poznata vrijednost s desne strane, "
            "udaljenost do svakog od tih rubova, relativan položaj točke unutar praznine te linearna "
            "procjena dobivena spajanjem dvaju rubova. Preostalih pet opisuje vrijeme mjerenja: "
            "normaliziran položaj u nizu te sat u danu i dan u godini, svaki prikazan sinusnom i "
            "kosinusnom komponentom.",
            "Značajke o praznini presudne su za razumijevanje rezultata. Bez njih model vidi samo kada "
            "je mjerenje obavljeno, ali ne i kolika je temperatura bila neposredno prije i poslije "
            "praznine, iako upravo taj podatak nosi gotovo svu informaciju u ovako gusto uzorkovanom "
            "nizu. Pri računanju značajki za poznatu točku ta se točka izostavlja iz vlastitog "
            "susjedstva, kako model ne bi mogao naučiti odgovor iz same sebe.",
            "Sve metode strojnog učenja uče odstupanje od linearne procjene, a ne samu temperaturu. "
            "Model dakle predviđa koliko linearna procjena griješi, a konačna vrijednost dobiva se "
            "zbrajanjem te procjene i naučenog odstupanja. Takva postavka daje modelu dobru polaznu "
            "točku: ako ne nauči ništa korisno, rezultat ostaje jednak linearnoj interpolaciji umjesto "
            "da bude znatno lošiji. Na kraju se svaka procjena ograničava na raspon poznatih "
            "temperatura u nizu, čime se sprječavaju vrijednosti izvan smislenog opsega.",
        ],
        oznaka="implementacija: znacajke",
    )

    # ---------------- IMPLEMENTACIJA: KNN ------------------------------------
    zamijeni(
        doc,
        "KNN metoda implementirana je tako da za svaku nedostajuću vrijednost traži najsličnije poznate "
        "uzorke u vremenskom nizu. Sličnost se računa pomoću značajki kao što su položaj u nizu, sat u "
        "danu i dan u godini. Nakon pronalaska k najbližih poznatih uzoraka, nedostajuća temperatura "
        "procjenjuje se kao prosjek njihovih temperatura. Napredna KNN metoda proširuje osnovni KNN "
        "pristup boljom pripremom ulaznih značajki. Položaj mjerenja u nizu najprije se normalizira "
        "kako njegova brojčana vrijednost ne bi imala prevelik utjecaj u odnosu na ostale značajke. "
        "Osim toga, sat u danu i dan u godini ne promatraju se kao obični linearni brojevi, nego se "
        "pretvaraju u sinusnu i kosinusnu komponentu. To je važno zato što vrijeme ima cikličku "
        "prirodu. Primjerice, sati 23:00 i 00:00 vremenski su vrlo blizu, ali ako se promatraju samo "
        "kao brojevi 23 i 0, model ih može pogrešno shvatiti kao jako udaljene vrijednosti. Korištenjem "
        "sinusne i kosinusne komponente takvi se vremenski podaci prikazuju kružno, pa se kraj jednog "
        "ciklusa prirodno povezuje s početkom sljedećeg. Na isti način može se opisati i dan u godini, "
        "jer se kraj godine nadovezuje na početak nove godine. Time KNN metoda može bolje prepoznati "
        "slične vremenske trenutke, primjerice slične sate u danu ili slična razdoblja godine. Napredna "
        "verzija dodatno koristi izračun u kojem sličniji poznati uzorci imaju veću važnost pri "
        "procjeni nedostajuće temperature, dok manje slični uzorci imaju manju važnost.",
        "Obje KNN izvedbe implementirane su zasebno jer rješavaju različit problem. Osnovni KNN za "
        "svaku prazninu binarnim traženjem pronalazi najbližu poznatu točku lijevo i najbližu desno te "
        "ih ponderira inverznom udaljenošću. Uz jednog susjeda sa svake strane taj je izračun "
        "matematički istovjetan linearnoj interpolaciji, jer se omjer težina 1/d₁ i 1/d₂ svodi na "
        "d₂/(d₁+d₂). Osnovni KNN time nije suparnik linearnoj interpolaciji nego njezino poopćenje, pa "
        "u tablicama rezultata daje istu vrijednost. Sat u danu i dan u godini izostavljeni su iz mjere "
        "udaljenosti jer na sedmodnevnom prozoru dan u godini poprima samo osam različitih vrijednosti, "
        "a mjerenja su pokazala da njihovo uključivanje rezultat pogoršava.",
        "implementacija: KNN prvi dio",
    )

    par_knn = nadi_pocetak(doc, "Obje KNN izvedbe implementirane su zasebno")
    if par_knn is not None and nadi_pocetak(doc, "Napredni KNN ne traži susjede po položaju") is None:
        klon_iza(
            tijelo, par_knn,
            "Napredni KNN ne traži susjede po položaju nego po sličnosti situacije u kojoj se praznina "
            "nalazi. Svaka se točka opisuje s pet veličina: relativnim položajem unutar praznine, "
            "logaritmiranim udaljenostima do lijevog i desnog poznatog ruba te sinusnom i kosinusnom "
            "komponentom sata u danu. Logaritam se koristi zato što udaljenosti do rubova variraju od "
            "jednog do nekoliko stotina uzoraka, pa bi bez njega duge praznine potpuno prevladale u "
            "mjeri udaljenosti. Za svaku nedostajuću vrijednost pronalazi se dvanaest poznatih točaka s "
            "najsličnijim opisom, a njihova odstupanja od linearne procjene ponderiraju se inverznom "
            "udaljenošću u tom prostoru. Dobiveni prosjek dodaje se linearnoj procjeni, pa metoda "
            "zapravo uči u kojim se situacijama linearna procjena sustavno vara.",
        )

    # ---------------- IMPLEMENTACIJA: stablo odlucivanja ---------------------
    zamijeni(
        doc,
        "Decision Tree metoda implementirana je kao stablo čvorova u kojem se podaci postupno dijele "
        "prema zadanim pravilima. Svaki unutarnji čvor sadrži pravilo podjele, odnosno odabranu "
        "značajku i prag prema kojem se uzorak usmjerava u lijevu ili desnu granu stabla. Na taj način "
        "se poznati uzorci postupno grupiraju prema sličnosti. Završni čvorovi stabla nazivaju se "
        "listovi, a u njima se nalazi konačna vrijednost predikcije, najčešće prosječna temperatura "
        "uzoraka koji su došli do tog lista. Izgradnja stabla temelji se na traženju podjele koja "
        "najbolje razdvaja poznate uzorke prema temperaturi. Za svaku moguću značajku i prag računa se "
        "pogreška podjele, a odabire se ona koja daje najmanju pogrešku. Postupak se ponavlja sve dok "
        "se ne dosegne najveća dopuštena dubina stabla ili dok u čvoru ne ostane premalo uzoraka za "
        "daljnju podjelu. Kod procjene nedostajuće vrijednosti uzorak prolazi kroz naučeno stablo, pri "
        "čemu se na svakom čvoru provjerava zadano pravilo. Kada uzorak dođe do lista, vrijednost tog "
        "lista koristi se kao procjena nedostajuće temperature.",
        "Decision Tree metoda implementirana je kao stablo čvorova u kojem se podaci postupno dijele "
        "prema zadanim pravilima. Svaki unutarnji čvor sadrži pravilo podjele, odnosno odabranu "
        "značajku i prag prema kojem se uzorak usmjerava u lijevu ili desnu granu stabla. Završni "
        "čvorovi nazivaju se listovi i u njima se nalazi konačna vrijednost predikcije. Izgradnja "
        "stabla temelji se na traženju podjele koja najbolje razdvaja poznate uzorke. Za svaku značajku "
        "uzorci se najprije sortiraju po njezinoj vrijednosti, nakon čega se jednim prolaskom kroz "
        "sortirani niz i uz pomoć kumulativnih zbrojeva ocjenjuju svi mogući pragovi, a odabire se onaj "
        "s najmanjom kvadratnom pogreškom. Postupak se ponavlja dok se ne dosegne dubina osam ili dok u "
        "čvoru ne ostane manje od četiri uzorka. Stablo pritom ne uči samu temperaturu nego njezino "
        "odstupanje od linearne procjene, uz jedanaest ranije opisanih značajki. Kod procjene "
        "nedostajuće vrijednosti uzorak prolazi kroz naučeno stablo do odgovarajućeg lista, a vrijednost "
        "tog lista dodaje se linearnoj procjeni. Značajke se računaju unaprijed za cijeli niz i "
        "pohranjuju u matricu, pa se isti izračun ne ponavlja u svakom čvoru.",
        "implementacija: stablo odlucivanja",
    )

    # ---------------- IMPLEMENTACIJA: slucajna suma --------------------------
    zamijeni(
        doc,
        "Random Forest metoda proširuje ideju stabla odlučivanja tako da koristi više stabala umjesto "
        "samo jednog. Svako stablo trenira se na nasumično odabranom skupu poznatih uzoraka s "
        "ponavljanjem, zbog čega pojedina stabla nisu potpuno jednaka i mogu dati različite procjene za "
        "istu nedostajuću vrijednost. Nakon izgradnje modela, svako stablo daje vlastitu procjenu, a "
        "konačna imputirana temperatura dobiva se kao prosjek svih dobivenih procjena. U kontekstu ovog "
        "rada Random Forest koristi vremenske značajke kako bi naučio odnos između položaja mjerenja, "
        "vremena i temperature. Prednost metode je veća stabilnost u odnosu na jedno stablo, dok su "
        "ograničenja veća složenost i teže objašnjenje pojedinačne odluke.",
        "Random Forest metoda proširuje ideju stabla odlučivanja tako da koristi 24 stabla umjesto "
        "jednog. Svako se stablo uči na nasumično odabranom skupu poznatih uzoraka s ponavljanjem, zbog "
        "čega pojedina stabla nisu jednaka i mogu dati različite procjene za istu nedostajuću "
        "vrijednost. Drugi izvor raznolikosti je odabir značajki: u svakom se čvoru razmatra nasumičnih "
        "sedam od jedanaest dostupnih značajki umjesto svih. Bez tog ograničenja sva bi stabla birala "
        "vrlo slične podjele, pa bi njihovo prosječenje donijelo malo. Stabla u šumi smiju biti dublja "
        "od pojedinačnog stabla, do dubine deset, jer se njihova sklonost pretjeranom prilagođavanju "
        "podacima djelomično poništava prosječenjem, uz uvjet da svaki list sadrži najmanje četiri "
        "uzorka. Kao i pojedinačno stablo, šuma uči odstupanje od linearne procjene, pa se konačna "
        "imputirana temperatura dobiva zbrajanjem linearne procjene i prosjeka odstupanja svih stabala. "
        "Prednost metode je veća stabilnost u odnosu na jedno stablo, dok su ograničenja veća složenost "
        "i teže objašnjenje pojedinačne odluke.",
        "implementacija: slucajna suma",
    )

    # ---------------- IMPLEMENTACIJA: neuronska mreza ------------------------
    umetni_odjeljak(
        doc,
        sidro=nadi_pocetak(doc, "Random Forest metoda proširuje ideju stabla odlučivanja tako da "
                                "koristi 24 stabla"),
        naslov_uzor=h_impl,
        tijelo_uzor=tijelo,
        naslov="Implementacija neuronske mreže",
        odlomci=[
            "Neuronska mreža implementirana je izravno u programskom jeziku C, bez vanjskih biblioteka, "
            "kako bi ostala usporediva s ostatkom sustava. Riječ je o višeslojnom perceptronu s "
            "jedanaest ulaza, dva skrivena sloja od 24 i 12 neurona te jednim izlazom. U skrivenim "
            "slojevima kao aktivacijska funkcija koristi se tangens hiperbolni, dok je izlaz linearan "
            "jer se predviđa realan broj, a ne razred.",
            "Prije učenja svaka se ulazna značajka standardizira oduzimanjem prosjeka i dijeljenjem "
            "standardnom devijacijom, pri čemu se te veličine računaju isključivo na poznatim točkama "
            "kako podaci o uklonjenim vrijednostima ne bi procurili u model. Izlazna veličina, "
            "odstupanje od linearne procjene, također se skalira jer je reda nekoliko desetinki "
            "stupnja, a mreža stabilnije uči na veličinama reda jedan. Početne težine postavljaju se "
            "Xavierovom inicijalizacijom, a težine izlaznog sloja dodatno se umanjuju kako bi mreža na "
            "početku učenja predviđala odstupanje blizu nule. Polazna točka učenja time je čista "
            "linearna interpolacija, pa je mreža može popraviti, ali teško i pogoršati.",
            "Učenje se provodi kroz 200 prolazaka nad nasumično promiješanim poznatim uzorcima, u "
            "serijama po 32 uzorka. Gradijenti se računaju propagacijom pogreške unatrag, a težine se "
            "osvježavaju Adam optimizatorom. [16] Stopa učenja počinje na 0,01 i kosinusno se gasi kroz "
            "prolaske, čime se završna faza učenja stabilizira. Nakon učenja mreža za svaku nedostajuću "
            "točku predviđa odstupanje, koje se dodaje linearnoj procjeni i ograničava na raspon "
            "poznatih temperatura.",
        ],
        oznaka="implementacija: neuronska mreza",
    )

    # ---------------- LITERATURA ---------------------------------------------
    zadnja_lit = None
    for par in doc.paragraphs:
        if par.style.name == "Literatura-radnja":
            zadnja_lit = par
    if zadnja_lit is None:
        promasaji.append("literatura")
    elif nadi_pocetak(doc, "Goodfellow, I.") is None:
        p = klon_iza(zadnja_lit, zadnja_lit,
                     "Goodfellow, I.; Bengio, Y.; Courville, A.: “Deep Learning”, MIT Press, "
                     "Cambridge, MA, 2016., s Interneta, https://www.deeplearningbook.org/.")
        klon_iza(zadnja_lit, p,
                 "Kingma, D. P.; Ba, J.: “Adam: A Method for Stochastic Optimization”, 3rd "
                 "International Conference on Learning Representations (ICLR), San Diego, CA, 2015., "
                 "s Interneta, https://arxiv.org/abs/1412.6980.")

    # ---------------- POPIS KRATICA ------------------------------------------
    t_krat = doc.tables[-1]
    postojeci = {t_krat.rows[i].cells[0].text.strip() for i in range(1, len(t_krat.rows))}
    nove = [("MLP", "Multilayer Perceptron; višeslojni perceptron, ovdje korištena neuronska mreža."),
            ("Adam", "Adaptive Moment Estimation; postupak osvježavanja težina pri učenju mreže."),
            ("rezidual", "Odstupanje stvarne vrijednosti od linearne procjene.")]
    nove = [r for r in nove if r[0] not in postojeci]
    if nove:
        staro_n = len(t_krat.rows) - 1
        podesi_broj_redaka(t_krat, staro_n + len(nove))
        for i, (a, b) in enumerate(nove):
            upisi_celiju(t_krat.rows[staro_n + 1 + i].cells[0], a)
            upisi_celiju(t_krat.rows[staro_n + 1 + i].cells[1], b)
        print(f"  popis kratica dopunjen: {len(nove)} novih redaka")

    # ---------------- SAZETAK, ABSTRACT, KLJUCNE RIJECI ----------------------
    zamijeni(
        doc,
        "U radu se ispituje koliko različiti postupci mogu pouzdano rekonstruirati nedostajuće "
        "temperature u vremenskom nizu. Eksperimenti su provedeni na temperaturnoj komponenti skupa "
        "Jena Climate, pri čemu su poznate vrijednosti namjerno uklanjane prema scenarijima random, "
        "block, block_start, block_middle i block_end. Uspoređeni su forward fill, linearna, vremenska, "
        "kubna i spline interpolacija te napredni KNN, Decision Tree i Random Forest. Missing rate "
        "mijenjan je od 10 % do 80 %, a pogreška je računata samo na uklonjenim pozicijama pomoću MAE, "
        "RMSE i R². U cjelokupnoj usporedbi linearna interpolacija pokazala se najstabilnijom, dok je "
        "kubna interpolacija ostvarila najbolje rezultate u random scenariju pri manjim udjelima "
        "nedostajućih vrijednosti.",
        "U radu se ispituje koliko različiti postupci mogu pouzdano rekonstruirati nedostajuće "
        "temperature u vremenskom nizu. Eksperimenti su provedeni na temperaturnoj komponenti skupa "
        "Jena Climate, pri čemu su poznate vrijednosti namjerno uklanjane prema scenarijima random, "
        "block, block_start, block_middle i block_end. Uspoređeno je jedanaest metoda: forward fill, "
        "linearna, vremenska, kubna i spline interpolacija, pomični prosjek, osnovni i napredni KNN, "
        "Decision Tree, Random Forest te neuronska mreža. Missing rate mijenjan je od 10 % do 80 %, a "
        "svaka je kombinacija ponovljena nad 20 neovisnih tjednih prozora kako zaključci ne bi ovisili "
        "o jednom razdoblju. Pogreška je računata samo na uklonjenim pozicijama pomoću MAE, RMSE i R². "
        "U cjelokupnoj usporedbi linearna interpolacija ostala je najstabilnija, ali su razlike prema "
        "najboljim metodama strojnog učenja manje od 0,08 °C, a na kontinuiranim blokovima prednost "
        "prelazi na njihovu stranu.",
        "sazetak",
    )

    zamijeni(
        doc,
        "This thesis examines how reliably different methods can reconstruct missing temperature values "
        "in a time series. Experiments were carried out on the temperature component of the Jena "
        "Climate Dataset, with known values deliberately removed using the random, block, block_start, "
        "block_middle, and block_end scenarios. The compared methods were forward fill, linear, time, "
        "cubic and spline interpolation, advanced KNN, Decision Tree, and Random Forest. Missing rates "
        "ranged from 10% to 80%, and MAE, RMSE, and R² were calculated only at the removed positions. "
        "Across the complete set of experiments, linear interpolation was the most stable approach, "
        "while cubic interpolation achieved the best results in the random scenario at lower missing "
        "rates.",
        "This thesis examines how reliably different methods can reconstruct missing temperature values "
        "in a time series. Experiments were carried out on the temperature component of the Jena "
        "Climate Dataset, with known values deliberately removed using the random, block, block_start, "
        "block_middle, and block_end scenarios. Eleven methods were compared: forward fill, linear, "
        "time, cubic and spline interpolation, moving average, basic and advanced KNN, Decision Tree, "
        "Random Forest, and a neural network. Missing rates ranged from 10% to 80%, and every "
        "combination was repeated over 20 independent weekly windows so that the conclusions would not "
        "depend on a single period. MAE, RMSE, and R² were calculated only at the removed positions. "
        "Across the complete set of experiments linear interpolation remained the most stable approach, "
        "but its margin over the best machine learning methods is below 0.08 °C, and on continuous "
        "blocks the advantage shifts to them.",
        "abstract",
    )

    zamijeni(
        doc,
        "Interpolacija podataka, imputacija, vremenski nizovi, strojno učenje, KNN, Random Forest",
        "Interpolacija podataka, imputacija, vremenski nizovi, strojno učenje, KNN, Random Forest, "
        "neuronska mreža",
        "kljucne rijeci",
    )
    zamijeni(
        doc,
        "Data interpolation, imputation, time series, machine learning, KNN, Random Forest",
        "Data interpolation, imputation, time series, machine learning, KNN, Random Forest, "
        "neural network",
        "keywords",
    )

    doc.save(str(DOK))
    print(f"\nzamijenjeno odlomaka: {brojac['zamjena']}, umetnuto novih: {brojac['umetnuto']}")
    if promasaji:
        print(f"NEPRONADENO ({len(promasaji)}):")
        for p in promasaji:
            print(f"   - {p}")
    else:
        print("sve ciljane izmjene provedene")
    print(f"spremljeno: {DOK.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
