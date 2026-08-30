#!/usr/bin/env python3
"""
Provjerava azurirani rad: strukturu, ocuvanost stilova i tocnost svake brojke.

Svaka brojcana celija u MAE tablicama usporeduje se s experiment_results.csv, a
tekst se pretrazuje na zaostale vrijednosti iz stare verzije eksperimenta.
"""

from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

import docx
import pandas as pd
from docx.oxml.ns import qn

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
STARI = ROOT / "docs" / "Diplomski-Toni_Jakelic_20.8_1.docx"
NOVI = ROOT / "docs" / "Diplomski-Toni_Jakelic_azurirano.docx"

METODE = {"forward_fill": "forward_fill", "linear": "linear_interpolation",
          "time": "time_interpolation", "cubic": "cubic_interpolation",
          "spline": "spline_interpolation", "moving_average": "moving_average",
          "knn": "knn", "knn_upgraded": "knn_upgraded",
          "decision_tree": "decision_tree", "random_forest": "random_forest",
          "neural_net": "neural_net"}
T_SCEN = {"random": 2, "block": 3, "block_start": 4, "block_middle": 5, "block_end": 6}
RATES = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8]

# Vrijednosti koje su postojale samo u staroj verziji eksperimenta.
ZAOSTALO = ["0,5855", "0,4353", "320 eksperimenata", "288 zapisa", "žmetoda",
            "27 pobjeda", "0,0919", "0,6705", "1,4370", "1,5937", "0,0406",
            "osam metoda", "1,4954", "1,8691", "198 ima R²", "0,2849", "0,3538",
            "1,1335", "0,3654", "0,0502", "3,4677", "3,3376",
            "pomičnog prosjeka nije uključena", "pet klasičnih metoda",
            "nije moguća jer se u završnom CSV-u"]

# Pojmovi koji moraju postojati nakon dopune poglavlja.
OBAVEZNO = ["višeslojnom perceptronu", "Adam optimizatorom", "Xavierovom inicijalizacijom",
            "jedanaest značajki", "odstupanje od linearne procjene", "24 stabla",
            "šest klasičnih metoda", "Goodfellow", "Kingma", "MLP"]


def main() -> None:
    stari = docx.Document(str(STARI))
    novi = docx.Document(str(NOVI))
    greske = 0

    print("=" * 88)
    print("1. STRUKTURA: stari naspram novog")
    print("=" * 88)
    for ime, d in (("stari", stari), ("novi", novi)):
        rijeci = sum(len(p.text.split()) for p in d.paragraphs)
        print(f"  {ime:<7} odlomaka {len(d.paragraphs):>4} | tablica {len(d.tables):>3} "
              f"| slika {len(d.inline_shapes):>3} | rijeci {rijeci:>6}")

    c_stari = Counter(p.style.name for p in stari.paragraphs)
    c_novi = Counter(p.style.name for p in novi.paragraphs)
    razlike = {k: (c_stari.get(k, 0), c_novi.get(k, 0))
               for k in set(c_stari) | set(c_novi) if c_stari.get(k) != c_novi.get(k)}
    print("\n  promjene u broju odlomaka po stilu (stari -> novi):")
    for k, (a, n) in sorted(razlike.items()):
        print(f"    {k:<24}{a:>4} -> {n:<4} ({n - a:+d})")
    # Dodana poglavlja smiju povecati samo ove stilove.
    dopusteni = {"Heading 2", "Heading 3", "Normal", "Literatura-radnja"}
    for k, (a, n) in razlike.items():
        if k not in dopusteni or n < a:
            print(f"  GRESKA: neocekivana promjena stila {k}: {a} -> {n}")
            greske += 1
    print(f"  stilovi tablica: {sorted({t.style.name for t in novi.tables})}")

    print()
    print("=" * 88)
    print("1b. NOVA POGLAVLJA")
    print("=" * 88)
    ocekivani = {"Značajke za metode strojnog učenja": "Heading 2",
                 "Implementacija neuronske mreže": "Heading 2",
                 "Neuronska mreža": "Heading 3"}
    nadeni = {p.text.strip(): p.style.name for p in novi.paragraphs
              if p.text.strip() in ocekivani}
    for naslov, stil in ocekivani.items():
        stvarni = nadeni.get(naslov)
        if stvarni == stil:
            print(f"  {naslov:<38} {stvarni}")
        else:
            print(f"  GRESKA: {naslov!r} ima stil {stvarni!r}, ocekivano {stil!r}")
            greske += 1

    print()
    print("=" * 88)
    print("2. NATPISI I POLJA ZA AUTOMATSKO NUMERIRANJE")
    print("=" * 88)
    polja_stari = sum(1 for p in stari.paragraphs for _ in p._element.iter(qn("w:instrText")))
    polja_novi = sum(1 for p in novi.paragraphs for _ in p._element.iter(qn("w:instrText")))
    print(f"  polja u starom: {polja_stari}   u novom: {polja_novi}")
    if polja_stari != polja_novi:
        print("  GRESKA: broj polja se promijenio")
        greske += 1
    natpisi = [p.text.strip() for p in novi.paragraphs
               if p.style.name == "Caption" and p.text.strip().startswith("Tablica")]
    for n in natpisi:
        print(f"  {n}")

    print()
    print("=" * 88)
    print("3. TABLICE: svaka celija naspram experiment_results.csv")
    print("=" * 88)
    res = pd.read_csv(ROOT / "results" / "experiment_results.csv")
    res["missing_rate"] = res.missing_rate.round(2)
    provjereno = 0
    for scen, idx in T_SCEN.items():
        tbl = novi.tables[idx]
        zag = [c.text.strip() for c in tbl.rows[0].cells]
        if zag[0] != "metoda":
            print(f"  GRESKA {scen}: zaglavlje je {zag[0]!r}")
            greske += 1
        for row in tbl.rows[1:]:
            celije = [c.text.strip() for c in row.cells]
            puno = METODE.get(celije[0])
            if puno is None:
                print(f"  GRESKA {scen}: nepoznata metoda {celije[0]!r}")
                greske += 1
                continue
            for j, rate in enumerate(RATES, start=1):
                d = res[(res.scenario == scen) & (res.method == puno)
                        & (res.missing_rate == rate)]
                ocekivano = f"{float(d.mae.iloc[0]):.4f}".replace(".", ",")
                if celije[j] != ocekivano:
                    print(f"  GRESKA {scen}/{celije[0]}/{rate}: "
                          f"u radu {celije[j]!r}, u CSV-u {ocekivano!r}")
                    greske += 1
                provjereno += 1
        print(f"  {scen:<14}{len(tbl.rows) - 1:>3} metoda x 8 rateova  u redu")
    print(f"\n  provjereno {provjereno} celija")

    t1, t7 = novi.tables[1], novi.tables[7]
    print(f"  tablica 5-1: {len(t1.rows) - 1} redaka (ocekivano 40)")
    print(f"  tablica 5-7: {len(t7.rows) - 1} redaka (ocekivano 11)")
    if len(t1.rows) - 1 != 40 or len(t7.rows) - 1 != 11:
        greske += 1

    print()
    print("=" * 88)
    print("4. ZAOSTALE VRIJEDNOSTI IZ STARE VERZIJE")
    print("=" * 88)
    tekst = "\n".join(p.text for p in novi.paragraphs)
    nasao = False
    for z in ZAOSTALO:
        if z in tekst:
            recenice = [s.strip() for s in re.split(r"(?<=\.)\s", tekst) if z in s]
            print(f"  JOS PRISUTNO {z!r}:")
            for s in recenice[:2]:
                print(f"      {s[:150]}")
            nasao = True
            greske += 1
    if not nasao:
        print("  nijedna stara vrijednost nije pronadena u tekstu")

    print()
    print("=" * 88)
    print("5. SADRZAJ NOVIH POGLAVLJA")
    print("=" * 88)
    svetekst = tekst + "\n".join(c.text for t in novi.tables for r in t.rows for c in r.cells)
    for o in OBAVEZNO:
        if o in svetekst:
            print(f"  nadeno: {o}")
        else:
            print(f"  GRESKA: nedostaje {o!r}")
            greske += 1

    print()
    print("=" * 88)
    print(f"UKUPNO GRESAKA: {greske}")
    print("=" * 88)


if __name__ == "__main__":
    main()
