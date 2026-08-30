#!/usr/bin/env python3
"""
Azurira diplomski rad na aktualno stanje eksperimenta.

Mijenja sadrzaj tablica 5-1 do 5-7 i one odlomke cije brojke ili tvrdnje vise
ne odgovaraju rezultatima iz results/experiment_results.csv. Stilovi, natpisi
tablica i slika te Wordova polja za automatsko numeriranje ostaju netaknuti:
tablice se popunjavaju u mjestu, redovi se dodaju kopiranjem XML-a postojeceg
retka, a tekst se upisuje u prvi run odlomka kako bi se zadrzalo oblikovanje.

Original se ne mijenja, rezultat se sprema pod novim imenom.
"""

from __future__ import annotations

import copy
import json
import re
import sys
from pathlib import Path

import docx

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
ULAZ = ROOT / "docs" / "Diplomski-Toni_Jakelic_20.8_1.docx"
IZLAZ = ROOT / "docs" / "Diplomski-Toni_Jakelic_azurirano.docx"
BROJKE = ROOT / "results" / "_brojke_za_rad.json"

SCEN = ["random", "block", "block_start", "block_middle", "block_end"]
# Indeksi tablica u dokumentu (0-bazirano).
T_NAJBOLJA = 1
T_SCEN = {"random": 2, "block": 3, "block_start": 4, "block_middle": 5, "block_end": 6}
T_SAZETAK = 7

promasaji: list[str] = []
odradeno = 0


# --------------------------------------------------------------------------
# Pomocne funkcije za rad s dokumentom
# --------------------------------------------------------------------------
def norm(s: str) -> str:
    """Normalizira razmake radi pouzdanog usporedivanja teksta iz Worda."""
    return re.sub(r"\s+", " ", s.replace("\u00a0", " ")).strip()


def upisi_odlomak(par, tekst: str) -> None:
    """
    Upisuje tekst u odlomak zadrzavajuci oblikovanje prvog runa.

    Odlomak bez izravnih runova drzi tekst u polju (npr. natpisi tablica), pa bi
    ga dodavanje novog runa udvostrucilo. Takav se slucaj odbija.
    """
    if not par.runs:
        if par.text.strip():
            raise RuntimeError(f"odlomak s poljem, ne diram ga: {par.text[:60]!r}")
        par.add_run(tekst)
        return
    par.runs[0].text = tekst
    for r in par.runs[1:]:
        r._element.getparent().remove(r._element)


def zamijeni_odlomak(doc, stari: str, novi: str, oznaka: str) -> None:
    global odradeno
    cilj = norm(stari)
    for par in doc.paragraphs:
        if norm(par.text) == cilj:
            upisi_odlomak(par, novi)
            odradeno += 1
            return
    promasaji.append(oznaka)


def zamijeni_u_runu(doc, sadrzi: str, staro: str, novo: str, oznaka: str) -> None:
    """
    Mijenja dio teksta unutar natpisa, na razini w:t cvorova.

    Natpisi tablica sadrze Wordovo polje za automatsko numeriranje, pa im tekst
    nije u izravnim runovima odlomka i mora se traziti medu svim potomcima.
    """
    global odradeno
    from docx.oxml.ns import qn
    for par in doc.paragraphs:
        if sadrzi in norm(par.text):
            for t in par._element.iter(qn("w:t")):
                if t.text and staro in t.text:
                    t.text = t.text.replace(staro, novo)
                    odradeno += 1
                    return
    promasaji.append(oznaka)


def upisi_celiju(cell, tekst: str) -> None:
    """Upisuje tekst u celiju zadrzavajuci font i podebljanje postojeceg runa."""
    par = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    upisi_odlomak(par, tekst)


def podesi_broj_redaka(tbl, treba: int) -> None:
    """Dodaje ili uklanja retke tako da tablica ima zaglavlje i 'treba' redaka."""
    ukupno = treba + 1
    while len(tbl.rows) < ukupno:
        novi = copy.deepcopy(tbl.rows[-1]._tr)
        tbl.rows[-1]._tr.addnext(novi)
    while len(tbl.rows) > ukupno:
        tbl.rows[-1]._tr.getparent().remove(tbl.rows[-1]._tr)


def popuni_tablicu(tbl, redovi: list[list[str]], zaglavlje: list[str] | None = None) -> None:
    podesi_broj_redaka(tbl, len(redovi))
    if zaglavlje:
        for j, v in enumerate(zaglavlje):
            if j < len(tbl.rows[0].cells):
                upisi_celiju(tbl.rows[0].cells[j], v)
    for i, red in enumerate(redovi, start=1):
        celije = tbl.rows[i].cells
        for j, v in enumerate(red):
            if j < len(celije):
                upisi_celiju(celije[j], v)


# --------------------------------------------------------------------------
def main() -> None:
    b = json.loads(BROJKE.read_text(encoding="utf-8"))
    doc = docx.Document(str(ULAZ))
    S = b["scenariji"]
    G = b["globalno"]
    R = b["rekonstrukcije"]
    # Dokument brojeve do deset pise rijecima, pa se isto radi i za broj metoda.
    n_met = {11: "jedanaest", 12: "dvanaest"}.get(b["n_metoda"], str(b["n_metoda"]))

    def pr(scen: str, rate: str, kljuc: str) -> str:
        return S[scen]["po_rateu"][rate][kljuc]

    # ---------------- tablice ------------------------------------------------
    popuni_tablicu(doc.tables[T_NAJBOLJA], b["tablica_najbolja"])
    for scen in SCEN:
        popuni_tablicu(doc.tables[T_SCEN[scen]], b["tablice_mae"][scen],
                       zaglavlje=["metoda", "10 %", "20 %", "30 %", "40 %",
                                  "50 %", "60 %", "70 %", "80 %"])
    popuni_tablicu(doc.tables[T_SAZETAK], b["tablica_sazetak"])
    print(f"tablice popunjene; zaglavlje 'žmetoda' ispravljeno u 'metoda'")

    # ---------------- natpis tablice 5-2: uskladi zapis postotka -------------
    zamijeni_u_runu(doc, "MAE vrijednosti za scenarij random missing",
                    "od 10% do 80%", "od 10 % do 80 %", "natpis tablice 5-2")

    # ---------------- broj uklonjenih vrijednosti ----------------------------
    zamijeni_odlomak(
        doc,
        "Testirani su missing rateovi od 10 %, 20 %, 30 %, 40 %, 50 %, 60 %, 70 % i 80 %. "
        "Manji postotak ostavlja više poznatih vrijednosti i olakšava rekonstrukciju, dok "
        "veći postotak smanjuje broj poznatih uzoraka i povećava težinu problema. "
        "Proširenjem do 80 % provjerava se koliko su metode stabilne u izrazito nepovoljnim "
        "uvjetima, kada je većina vrijednosti u nizu umjetno uklonjena. Na Jena skupu od 288 "
        "zapisa to znači približno 29, 58, 86, 115, 144, 173, 202 i 230 uklonjenih vrijednosti. "
        "Broj stvarno uklonjenih i evaluiranih vrijednosti zapisan je u stupcima "
        "number_of_missing_values i number_of_evaluated_values u glavnoj CSV datoteci rezultata.",
        "Testirani su missing rateovi od 10 %, 20 %, 30 %, 40 %, 50 %, 60 %, 70 % i 80 %. "
        "Manji postotak ostavlja više poznatih vrijednosti i olakšava rekonstrukciju, dok "
        "veći postotak smanjuje broj poznatih uzoraka i povećava težinu problema. "
        "Proširenjem do 80 % provjerava se koliko su metode stabilne u izrazito nepovoljnim "
        "uvjetima, kada je većina vrijednosti u nizu umjetno uklonjena. Na tjednom prozoru od "
        "1008 zapisa to znači redom 101, 202, 302, 403, 504, 605, 706 i 806 uklonjenih "
        "vrijednosti. Broj stvarno uklonjenih i evaluiranih vrijednosti zapisan je u stupcima "
        "number_of_missing_values i number_of_evaluated_values u glavnoj CSV datoteci rezultata.",
        "broj uklonjenih vrijednosti",
    )

    # ---------------- uvod u poglavlje REZULTATI -----------------------------
    zamijeni_odlomak(
        doc,
        "U ovom poglavlju prikazani su rezultati eksperimenata nad temperaturnim vremenskim "
        "nizom iz Jena Climate skupa podataka. Provedeno je ukupno 320 eksperimenata, odnosno "
        "sve kombinacije pet scenarija, osam missing rateova od 10 % do 80 % i osam metoda "
        "imputacije. Cilj nije samo odrediti koja metoda ima najmanju pogrešku, nego pokazati "
        "u kojim uvjetima pojedine metode zadržavaju dobru rekonstrukciju, a u kojim uvjetima "
        "njihove pogreške rastu.",
        "U ovom poglavlju prikazani su rezultati eksperimenata nad temperaturnim vremenskim "
        "nizom iz Jena Climate skupa podataka. Provedeno je ukupno "
        f"{b['n_kombinacija'] * b['n_metoda']} kombinacija, odnosno svih pet scenarija, osam "
        f"missing rateova od 10 % do 80 % i {n_met} metoda imputacije. Svaka je "
        f"kombinacija ponovljena nad {b['n_ponavljanja']} međusobno neovisnih tjednih prozora "
        "izdvojenih iz različitih dijelova godine, pa ukupan broj pojedinačnih izvođenja iznosi "
        f"{b['n_izvodenja']}. Cilj nije samo odrediti koja metoda ima najmanju pogrešku, nego "
        "pokazati u kojim uvjetima pojedine metode zadržavaju dobru rekonstrukciju, a u kojim "
        "uvjetima njihove pogreške rastu.",
        "uvod u REZULTATE",
    )

    zamijeni_odlomak(
        doc,
        "Rezultati iz datoteke experiment_results.csv analizirani su kroz MAE, RMSE i R². Za "
        "MAE i RMSE manje vrijednosti znače manju pogrešku rekonstrukcije, dok je kod R² "
        "poželjnija veća vrijednost. U glavnom dijelu rada zadržane su MAE tablice i sažetak "
        "najboljih metoda, a RMSE i R² prikazani su grafički kako bi usporedba ostala pregledna "
        "bez ponavljanja gotovo jednakih tablica.",
        "Rezultati iz datoteke experiment_results.csv analizirani su kroz MAE, RMSE i R². Za "
        "MAE i RMSE manje vrijednosti znače manju pogrešku rekonstrukcije, dok je kod R² "
        "poželjnija veća vrijednost. Svaka vrijednost u tablicama koje slijede prosjek je "
        f"{b['n_ponavljanja']} ponavljanja, čime zaključci prestaju ovisiti o jednom slučajno "
        "odabranom tjednu i jednoj slučajnoj maski. Linearna interpolacija, vremenska "
        "interpolacija i KNN daju nad ovim nizom brojčano istovjetne rezultate, pa se pri "
        "određivanju najbolje metode navode zajedno kao izjednačene umjesto da se proizvoljno "
        "odabere jedna od njih. U glavnom dijelu rada zadržane su MAE tablice i sažetak "
        "najboljih metoda, a RMSE i R² prikazani su grafički kako bi usporedba ostala pregledna "
        "bez ponavljanja gotovo jednakih tablica.",
        "drugi uvodni odlomak REZULTATA",
    )

    zamijeni_odlomak(
        doc,
        "Iz Tablice 5-1 vidi se da linearna interpolacija najčešće završava kao metoda s "
        "najmanjim MAE-om: najbolja je u 27 od 40 kombinacija scenarija i missing ratea. Kubna "
        "interpolacija preuzima prednost uglavnom u random scenariju pri manjim udjelima "
        "nedostajućih vrijednosti, dok se kod blokova rezultat više mijenja ovisno o njihovoj "
        "poziciji i duljini.",
        "Iz Tablice 5-1 vidi se da nijedna metoda ne dominira kroz sve uvjete. Linearna "
        "interpolacija, zajedno s vremenskom interpolacijom i KNN-om, ima najmanji MAE u 15 od "
        "40 kombinacija scenarija i missing ratea, ponajprije pri višim udjelima nedostajućih "
        "vrijednosti u random scenariju te u scenarijima block_start i block_end. Metode "
        "strojnog učenja najbolje su u 20 kombinacija, sve do jedne na kontinuiranim "
        "blokovima: napredni KNN u osam, neuronska mreža u šest, stablo odlučivanja u četiri i "
        "slučajna šuma u dvije. Kubnoj i spline interpolaciji pripada preostalih pet "
        "kombinacija, i to samo pri najnižim missing rateovima.",
        "komentar uz tablicu 5-1",
    )

    # ---------------- random -------------------------------------------------
    zamijeni_odlomak(
        doc,
        "Kubna interpolacija ima najmanji MAE pri 10 %, 20 % i 30 % nedostajućih vrijednosti, "
        "redom 0,0406, 0,0488 i 0,0448. Od 40 % nadalje prvo mjesto preuzima linearna "
        "interpolacija. Čak i pri 80 % ona zadržava MAE 0,0919, RMSE 0,1259 i R² 0,9964, pa se "
        "random scenarij pokazuje daleko manje zahtjevnim od scenarija s kontinuiranim blokovima.",
        "Spline interpolacija ima najmanji MAE pri 10 % i 20 % nedostajućih vrijednosti, s "
        f"{pr('random', '10', 'najbolji_mae')} i {pr('random', '20', 'najbolji_mae')}, a kubna "
        f"interpolacija pri 30 % i 40 %, s {pr('random', '30', 'najbolji_mae')} i "
        f"{pr('random', '40', 'najbolji_mae')}. Od 50 % nadalje prvo mjesto preuzima linearna "
        "interpolacija, zajedno s vremenskom interpolacijom i KNN-om. Čak i pri 80 % ona "
        f"zadržava MAE {S['random']['lin_80_mae']}, RMSE {S['random']['lin_80_rmse']} i R² "
        f"{S['random']['lin_80_r2']}, pa se random scenarij pokazuje daleko manje zahtjevnim od "
        "scenarija s kontinuiranim blokovima.",
        "random uvod",
    )

    zamijeni_odlomak(
        doc,
        "Usporedba MAE-a po metodama pokazuje jasnu prednost interpolacijskih postupaka u random "
        "scenariju. Linearna interpolacija ima najmanju pogrešku u pet od osam promatranih "
        "missing rateova, dok se kod nižih rateova ističe kubna interpolacija.",
        "Usporedba MAE-a po metodama pokazuje jasnu prednost interpolacijskih postupaka u random "
        "scenariju. Linearna interpolacija ima najmanju pogrešku u četiri od osam promatranih "
        "missing rateova, dok se pri nižim rateovima ističu kubna i spline interpolacija. Metode "
        "strojnog učenja drže se vrlo blizu, ali nijedna od njih ne preuzima prvo mjesto.",
        "random usporedba MAE",
    )

    zamijeni_odlomak(
        doc,
        "Pri najvećem testiranom udjelu od 80 % razlika među metodama i dalje je velika: linear "
        "završava na MAE 0,0919, dok knn doseže 0,6353.",
        "Pri najvećem testiranom udjelu od 80 % razlika među metodama i dalje je velika: linear "
        f"završava na MAE {S['random']['lin_80_mae']}, dok forward_fill doseže 0,5769.",
        "random 80 % raspon",
    )

    zamijeni_odlomak(
        doc,
        "S porastom missing ratea prosječni MAE svih metoda raste s 0,0791 pri 10 % na 0,2240 pri "
        "80 %. Unatoč tom rastu, linearna interpolacija pri 80 % ostaje na samo 0,0919, što "
        "potvrđuje da pojedinačno raspoređene praznine i dalje ostavljaju dovoljno lokalnih "
        "informacija za dobru rekonstrukciju.",
        f"S porastom missing ratea prosječni MAE svih metoda raste s {pr('random', '10', 'prosjek_mae')} "
        f"pri 10 % na {pr('random', '80', 'prosjek_mae')} pri 80 %. Unatoč tom rastu, linearna "
        f"interpolacija pri 80 % ostaje na samo {S['random']['lin_80_mae']}, što potvrđuje da "
        "pojedinačno raspoređene praznine i dalje ostavljaju dovoljno lokalnih informacija za "
        "dobru rekonstrukciju.",
        "random MAE kroz rateove",
    )

    zamijeni_odlomak(
        doc,
        "RMSE slijedi isti opći smjer kao MAE, ali jače kažnjava veća pojedinačna odstupanja. "
        "Prosjek svih metoda raste s 0,1020 na 0,3352 između 10 % i 80 % missing ratea, dok "
        "linearna interpolacija pri 80 % ostvaruje RMSE 0,1259.",
        "RMSE slijedi isti opći smjer kao MAE, ali jače kažnjava veća pojedinačna odstupanja. "
        f"Prosjek svih metoda raste s {pr('random', '10', 'prosjek_rmse')} na "
        f"{pr('random', '80', 'prosjek_rmse')} između 10 % i 80 % missing ratea, dok linearna "
        f"interpolacija pri 80 % ostvaruje RMSE {S['random']['lin_80_rmse']}.",
        "random RMSE",
    )

    zamijeni_odlomak(
        doc,
        "R² u random scenariju ostaje vrlo visok za najbolje metode i pri velikom broju uklonjenih "
        "vrijednosti. Prosjek svih metoda smanjuje se s 0,9959 pri 10 % na 0,9543 pri 80 %, a "
        "linearna interpolacija na 80 % postiže 0,9964. To znači da njezina rekonstrukcija i dalje "
        "vrlo dobro prati oblik originalnog temperaturnog niza.",
        "R² u random scenariju ostaje vrlo visok za najbolje metode i pri velikom broju uklonjenih "
        f"vrijednosti. Prosjek svih metoda smanjuje se s {pr('random', '10', 'prosjek_r2')} pri "
        f"10 % na {pr('random', '80', 'prosjek_r2')} pri 80 %, a linearna interpolacija na 80 % "
        f"postiže {S['random']['lin_80_r2']}. To znači da njezina rekonstrukcija i dalje vrlo "
        "dobro prati oblik originalnog temperaturnog niza.",
        "random R2",
    )

    # ---------------- block --------------------------------------------------
    zamijeni_odlomak(
        doc,
        "Linearna interpolacija ima najmanji MAE pri 10 %, 30 %, 40 %, 60 %, 70 % i 80 %. Iznimke "
        "su 20 %, gdje je najbolja cubic_interpolation s MAE 0,1946, te 50 %, gdje forward_fill "
        "ostvaruje 0,6913. Na 80 % linearna interpolacija završava s MAE 0,6705, RMSE 0,9274 i R² "
        "0,7914.",
        "Prvo mjesto u ovom scenariju pripada isključivo metodama strojnog učenja. Neuronska mreža "
        "najbolja je pri 10 %, 20 %, 30 % i 60 %, slučajna šuma pri 40 % i 50 %, a stablo "
        "odlučivanja pri 70 % i 80 %. Prednost je pritom vrlo mala: pri 10 % neuronska mreža "
        f"ostvaruje MAE {pr('block', '10', 'najbolji_mae')} naspram 2,1669 linearne interpolacije. "
        f"Na 80 % linearna interpolacija završava s MAE {S['block']['lin_80_mae']}, RMSE "
        f"{S['block']['lin_80_rmse']} i R² {S['block']['lin_80_r2']}.",
        "block uvod",
    )

    zamijeni_odlomak(
        doc,
        "Kod kontinuiranog bloka razlike među metodama postaju izraženije nego u random scenariju. "
        "Linearna interpolacija ima najmanji MAE u šest od osam rateova. Na 80 % njezin MAE iznosi "
        "0,6705, dok decision_tree u istoj postavci doseže 2,8560.",
        "Kod kontinuiranog bloka razlike među metodama postaju izraženije nego u random scenariju, "
        "ali se sve upotrebljive metode grupiraju vrlo blizu jedna drugoj. Linearna interpolacija i "
        "četiri metode strojnog učenja na 80 % razlikuju se za manje od 0,02 °C, dok kubna "
        "interpolacija u istoj postavci doseže 29,3598.",
        "block usporedba MAE",
    )

    zamijeni_odlomak(
        doc,
        "Prosječni MAE svih metoda raste s 0,5387 pri 10 % na 1,8690 pri 80 %. Takav porast pokazuje "
        "koliko kontinuirana praznina smanjuje količinu lokalnih informacija dostupnih za "
        "rekonstrukciju. Linearna interpolacija je i na 80 % najbolja s MAE 0,6705.",
        f"Prosječni MAE svih metoda raste s {pr('block', '10', 'prosjek_mae')} pri 10 % na "
        f"{pr('block', '80', 'prosjek_mae')} pri 80 %. Takav porast pokazuje koliko kontinuirana "
        "praznina smanjuje količinu lokalnih informacija dostupnih za rekonstrukciju, ali i koliko "
        "na prosjek utječu metode koje se na blokovima raspadaju. Najbolja metoda na 80 % je stablo "
        f"odlučivanja s MAE {pr('block', '80', 'najbolji_mae')}.",
        "block MAE kroz rateove",
    )

    zamijeni_odlomak(
        doc,
        "Kod RMSE-a razlika se dodatno povećava jer ova metrika snažnije reagira na veća odstupanja. "
        "Prosječna vrijednost raste s 0,5942 na 2,2323, dok linearna interpolacija pri 80 % zadržava "
        "RMSE 0,9274.",
        "Kod RMSE-a razlika se dodatno povećava jer ova metrika snažnije reagira na veća odstupanja. "
        f"Prosječna vrijednost raste s {pr('block', '10', 'prosjek_rmse')} na "
        f"{pr('block', '80', 'prosjek_rmse')}, dok linearna interpolacija pri 80 % zadržava RMSE "
        f"{S['block']['lin_80_rmse']}.",
        "block RMSE",
    )

    zamijeni_odlomak(
        doc,
        "Za R² se ne dobiva jednoličan trend kao kod MAE-a i RMSE-a. Prosjek je pri 10 % izrazito "
        "negativan (-18,6495), dok se pri 80 % približava nuli (-0,3851). Ipak, najbolja pojedinačna "
        "metoda pri 80 % je linearna interpolacija s R² = 0,7914, što pokazuje da prosjek snažno "
        "snižavaju slabije metode.",
        "Za R² se ne dobiva jednoličan trend kao kod MAE-a i RMSE-a. Prosjek pada s "
        f"{pr('block', '10', 'prosjek_r2')} pri 10 % na {pr('block', '80', 'prosjek_r2')} pri 80 %, "
        "no taj pad gotovo u cijelosti proizlazi iz kubne i spline interpolacije. Najbolja "
        "pojedinačna metoda pri 80 % je stablo odlučivanja s R² -0,5277, što pokazuje da prosjek "
        "snažno snižavaju slabije metode.",
        "block R2",
    )

    # ---------------- block_start --------------------------------------------
    zamijeni_odlomak(
        doc,
        "U block_start scenariju nedostajući dio nalazi se na početku niza, pa metode nemaju poznate "
        "vrijednosti prije praznine. Rezultati pri nižim rateovima nisu jednolični: na 20 % najmanji "
        "MAE ima forward_fill, a na 30 % cubic_interpolation. Od 40 % do 80 % najbolja je "
        "linear_interpolation, koja pri 80 % ostvaruje MAE 0,6939, RMSE 0,9373 i R² 0,7794.",
        "U block_start scenariju nedostajući dio nalazi se na početku niza, pa metode nemaju poznate "
        "vrijednosti prije praznine. Rezultati pri nižim rateovima nisu jednolični: na 10 % najmanji "
        f"MAE ima kubna interpolacija s {pr('block_start', '10', 'najbolji_mae')}, a na 20 % napredni "
        f"KNN s {pr('block_start', '20', 'najbolji_mae')}. Od 30 % do 80 % najbolja je linearna "
        f"interpolacija, koja pri 80 % ostvaruje MAE {S['block_start']['lin_80_mae']}, RMSE "
        f"{S['block_start']['lin_80_rmse']} i R² {S['block_start']['lin_80_r2']}.",
        "block_start uvod",
    )

    zamijeni_odlomak(
        doc,
        "Raspored MAE vrijednosti u block_start scenariju pokazuje da linearna interpolacija pobjeđuje "
        "u šest od osam rateova. Pri 80 % razlika je posebno vidljiva: linear ima MAE 0,6939, dok "
        "decision_tree završava na 3,0026.",
        "Raspored MAE vrijednosti u block_start scenariju pokazuje da linearna interpolacija pobjeđuje "
        "u šest od osam rateova. Pri 80 % razlika prema slabijim metodama posebno je vidljiva: linear "
        f"ima MAE {S['block_start']['lin_80_mae']}, dok spline interpolacija završava na 15,0115.",
        "block_start usporedba MAE",
    )

    zamijeni_odlomak(
        doc,
        "Prosječni MAE u ovom scenariju ne raste potpuno pravilno s missing rateom, ali se ukupno "
        "povećava s 0,8956 pri 10 % na 1,6005 pri 80 %. Na najvišem rateu linearna interpolacija "
        "ostaje najbolja s vrijednošću 0,6939.",
        "Prosječni MAE u ovom scenariju ne raste potpuno pravilno s missing rateom, ali se ukupno "
        f"povećava s {pr('block_start', '10', 'prosjek_mae')} pri 10 % na "
        f"{pr('block_start', '80', 'prosjek_mae')} pri 80 %. Na najvišem rateu linearna interpolacija "
        f"ostaje najbolja s vrijednošću {S['block_start']['lin_80_mae']}.",
        "block_start MAE kroz rateove",
    )

    zamijeni_odlomak(
        doc,
        "Sličan obrazac vidi se i kod RMSE-a. Prosjek svih metoda iznosi 1,0044 pri 10 % i 1,9898 pri "
        "80 %, dok linearna interpolacija na 80 % postiže znatno nižih 0,9373.",
        "Sličan obrazac vidi se i kod RMSE-a. Prosjek svih metoda iznosi "
        f"{pr('block_start', '10', 'prosjek_rmse')} pri 10 % i {pr('block_start', '80', 'prosjek_rmse')} "
        f"pri 80 %, dok linearna interpolacija na 80 % postiže znatno nižih "
        f"{S['block_start']['lin_80_rmse']}.",
        "block_start RMSE",
    )

    zamijeni_odlomak(
        doc,
        "R² dodatno pokazuje nestabilnost pojedinih metoda na rubnom bloku. Prosječna vrijednost "
        "mijenja se s -7,5930 pri 10 % na -0,2532 pri 80 %, pa prosjek nije dobar pokazatelj ponašanja "
        "najbolje metode. Linearna interpolacija na 80 % ostvaruje R² 0,7794.",
        "R² dodatno pokazuje nestabilnost pojedinih metoda na rubnom bloku. Prosječna vrijednost "
        f"mijenja se s {pr('block_start', '10', 'prosjek_r2')} pri 10 % na "
        f"{pr('block_start', '80', 'prosjek_r2')} pri 80 %, pa prosjek nije dobar pokazatelj ponašanja "
        f"najbolje metode. Linearna interpolacija na 80 % ostvaruje R² {S['block_start']['lin_80_r2']}.",
        "block_start R2",
    )

    # ---------------- block_middle -------------------------------------------
    zamijeni_odlomak(
        doc,
        "Pri 10 % najmanji MAE ima decision_tree (0,1081), a pri 20 % forward_fill (0,3384). Od 30 % do "
        "80 % najbolja je linear_interpolation. Na 80 % i ona osjetno griješi više nego pri nižim "
        "rateovima, s MAE 1,4370 i RMSE 1,5759.",
        f"Pri 10 % i 50 % najmanji MAE ima napredni KNN, s {pr('block_middle', '10', 'najbolji_mae')} i "
        f"{pr('block_middle', '50', 'najbolji_mae')}, pri 20 % i 40 % neuronska mreža, s "
        f"{pr('block_middle', '20', 'najbolji_mae')} i {pr('block_middle', '40', 'najbolji_mae')}, a pri "
        f"30 % i 80 % stablo odlučivanja, s {pr('block_middle', '30', 'najbolji_mae')} i "
        f"{pr('block_middle', '80', 'najbolji_mae')}. Napredni KNN najbolji je i pri 70 %. Linearna "
        f"interpolacija prvo mjesto zauzima jedino pri 60 %, s {pr('block_middle', '60', 'najbolji_mae')}, "
        f"a na 80 % ostvaruje MAE {S['block_middle']['lin_80_mae']} i RMSE "
        f"{S['block_middle']['lin_80_rmse']}.",
        "block_middle uvod",
    )

    zamijeni_odlomak(
        doc,
        "U šest od osam postavki block_middle scenarija najmanji MAE postiže linearna interpolacija. Pri "
        "80 % njezina pogreška iznosi 1,4370, dok cubic u istom slučaju raste na 5,3329.",
        "U sedam od osam postavki block_middle scenarija najmanji MAE postiže neka od metoda strojnog "
        "učenja, no njihova prednost pred linearnom interpolacijom nigdje ne prelazi 0,03 °C. Pri 80 % "
        f"pogreška linearne interpolacije iznosi {S['block_middle']['lin_80_mae']}, dok kubna u istom "
        "slučaju raste na 24,9233.",
        "block_middle usporedba MAE",
    )

    zamijeni_odlomak(
        doc,
        "Povećanjem količine uklonjenih podataka prosječni MAE svih metoda raste s 0,3797 pri 10 % na "
        "3,3376 pri 80 %. Linearna interpolacija na 80 % i dalje je najbolja, ali s MAE 1,4370 rezultat "
        "je osjetno slabiji nego u random scenariju.",
        "Povećanjem količine uklonjenih podataka prosječni MAE svih metoda raste s "
        f"{pr('block_middle', '10', 'prosjek_mae')} pri 10 % na {pr('block_middle', '80', 'prosjek_mae')} "
        f"pri 80 %. Linearna interpolacija na 80 % s MAE {S['block_middle']['lin_80_mae']} ostaje blizu "
        "najboljeg rezultata, ali je i taj rezultat osjetno slabiji nego u random scenariju.",
        "block_middle MAE kroz rateove",
    )

    zamijeni_odlomak(
        doc,
        "RMSE potvrđuje isti problem s velikim središnjim blokom: prosjek raste s 0,5028 na 3,8235. "
        "Najbolji rezultat na 80 % ponovno ima linearna interpolacija, s RMSE 1,5759.",
        "RMSE potvrđuje isti problem s velikim središnjim blokom: prosjek raste s "
        f"{pr('block_middle', '10', 'prosjek_rmse')} na {pr('block_middle', '80', 'prosjek_rmse')}. "
        f"Linearna interpolacija na 80 % ostvaruje RMSE {S['block_middle']['lin_80_rmse']}.",
        "block_middle RMSE",
    )

    zamijeni_odlomak(
        doc,
        "Prosječni R² u block_middle scenariju ostaje negativan kroz promatrane krajnje točke i mijenja "
        "se s -33,0389 pri 10 % na -4,2901 pri 80 %. Linearna interpolacija je pri 80 % znatno bolja od "
        "prosjeka, ali R² 0,2636 pokazuje da je rekonstrukcija velikog središnjeg bloka i dalje "
        "zahtjevna.",
        "Prosječni R² u block_middle scenariju ostaje negativan kroz promatrane krajnje točke i mijenja "
        f"se s {pr('block_middle', '10', 'prosjek_r2')} pri 10 % na "
        f"{pr('block_middle', '80', 'prosjek_r2')} pri 80 %. Linearna interpolacija je pri 80 % znatno "
        f"bolja od prosjeka, ali R² {S['block_middle']['lin_80_r2']} pokazuje da je rekonstrukcija "
        "velikog središnjeg bloka i dalje zahtjevna.",
        "block_middle R2",
    )

    # ---------------- block_end ----------------------------------------------
    zamijeni_odlomak(
        doc,
        "Najbolja metoda mijenja se više nego u prethodnim scenarijima: linear_interpolation pobjeđuje "
        "pri 10 %, 20 %, 70 % i 80 %, cubic_interpolation pri 30 % i 60 %, random_forest pri 40 %, a "
        "decision_tree pri 50 %. Pri 80 % prosječni MAE svih metoda doseže 3,4677, najveću vrijednost "
        "među promatranim scenarijima.",
        "Najbolja metoda izmjenjuje se između dva pristupa: napredni KNN pobjeđuje pri 10 %, 20 %, 50 % "
        "i 60 %, a linearna interpolacija pri 30 %, 40 %, 70 % i 80 %. Pri 80 % prosječni MAE svih "
        f"metoda doseže {pr('block_end', '80', 'prosjek_mae')}.",
        "block_end uvod",
    )

    zamijeni_odlomak(
        doc,
        "Kod block_end scenarija nema jedne metode koja dominira kroz cijeli raspon, iako linearna "
        "interpolacija ima najviše pobjeda, četiri od osam. Na 80 % njezin MAE iznosi 1,5937, dok cubic "
        "raste na 6,4665.",
        "Kod block_end scenarija nema jedne metode koja dominira kroz cijeli raspon; linearna "
        "interpolacija i napredni KNN dijele pobjede po četiri. Na 80 % MAE linearne interpolacije "
        f"iznosi {S['block_end']['lin_80_mae']}, dok spline raste na 19,9462.",
        "block_end usporedba MAE",
    )

    zamijeni_odlomak(
        doc,
        "Prosječni MAE se s 0,4220 pri 10 % povećava na 3,4677 pri 80 %, što je najveći završni prosjek "
        "među svim scenarijima. Linearna interpolacija i u tom najtežem slučaju ostaje najbolja s MAE "
        "1,5937.",
        f"Prosječni MAE se s {pr('block_end', '10', 'prosjek_mae')} pri 10 % povećava na "
        f"{pr('block_end', '80', 'prosjek_mae')} pri 80 %. Linearna interpolacija i u tom slučaju ostaje "
        f"najbolja s MAE {S['block_end']['lin_80_mae']}.",
        "block_end MAE kroz rateove",
    )

    zamijeni_odlomak(
        doc,
        "RMSE pokazuje još izraženiji rast pogreške na završnom bloku: prosjek ide s 0,5453 pri 10 % na "
        "3,8395 pri 80 %. Linearna interpolacija na 80 % postiže RMSE 1,7404, znatno manje od prosjeka "
        "ostalih metoda.",
        "RMSE pokazuje još izraženiji rast pogreške na završnom bloku: prosjek ide s "
        f"{pr('block_end', '10', 'prosjek_rmse')} pri 10 % na {pr('block_end', '80', 'prosjek_rmse')} pri "
        f"80 %. Linearna interpolacija na 80 % postiže RMSE {S['block_end']['lin_80_rmse']}, znatno manje "
        "od prosjeka ostalih metoda.",
        "block_end RMSE",
    )

    zamijeni_odlomak(
        doc,
        "Kod R² se block_end razlikuje od prethodna dva rubna/središnja slučaja jer se prosječna "
        "vrijednost s povećanjem ratea snažno pogoršava: od -0,6474 pri 10 % do -11,9435 pri 80 %. Čak i "
        "najbolja metoda na 80 %, linearna interpolacija, ima R² -1,0366, što potvrđuje da je završni "
        "blok te veličine vrlo teško rekonstruirati.",
        "Kod R² se prosječna vrijednost s povećanjem ratea snažno pogoršava: od "
        f"{pr('block_end', '10', 'prosjek_r2')} pri 10 % do {pr('block_end', '80', 'prosjek_r2')} pri "
        "80 %. Čak i najbolja metoda na 80 %, linearna interpolacija, ima R² "
        f"{S['block_end']['lin_80_r2']}, što potvrđuje da je završni blok te veličine vrlo teško "
        "rekonstruirati.",
        "block_end R2",
    )

    # ---------------- utjecaj missing ratea ----------------------------------
    zamijeni_odlomak(
        doc,
        "Kada se svi scenariji promatraju zajedno, veći missing rate uglavnom znači i veću pogrešku. "
        "Prosječni MAE raste s 0,4630 pri 10 % na 2,0998 pri 80 %, a prosječni RMSE s 0,5497 na 2,4441. "
        "Najmanji utjecaj povećanja ratea vidi se u random scenariju, gdje oko uklonjenih točaka često "
        "ostaju poznata mjerenja. Kontinuirani blokovi puno su osjetljiviji, posebno block_end i "
        "block_middle pri 70 % i 80 %. Na 80 % prosječni MAE iznosi 3,4677 za block_end i 3,3376 za "
        "block_middle. R² ne prati missing rate tako pravilno kao MAE i RMSE jer njegova vrijednost "
        "ovisi i o varijabilnosti stvarnih temperatura na maskiranim pozicijama. Negativni rezultati "
        "zato su česti u block scenarijima, osobito za KNN, forward fill te cubic/spline interpolaciju "
        "na rubnim blokovima. Od ukupno 320 eksperimenata, 198 ima R² manji od nule.",
        "Kada se svi scenariji promatraju zajedno, veći missing rate uglavnom znači i veću pogrešku. "
        f"Prosječni MAE raste s {G['mae_10']} pri 10 % na {G['mae_80']} pri 80 %, a prosječni RMSE s "
        f"{G['rmse_10']} na {G['rmse_80']}. Najmanji utjecaj povećanja ratea vidi se u random scenariju, "
        "gdje oko uklonjenih točaka često ostaju poznata mjerenja. Kontinuirani blokovi puno su "
        f"osjetljiviji: na 80 % prosječni MAE iznosi {pr('block', '80', 'prosjek_mae')} za block i "
        f"{pr('block_middle', '80', 'prosjek_mae')} za block_middle. R² ne prati missing rate tako "
        "pravilno kao MAE i RMSE jer njegova vrijednost ovisi i o varijabilnosti stvarnih temperatura na "
        "maskiranim pozicijama. Negativni rezultati zato su česti u block scenarijima, osobito za "
        "forward fill, pomični prosjek te kubnu i spline interpolaciju. Od ukupno "
        f"{G['r2_ukupno']} kombinacija, {G['r2_negativnih']} ima R² manji od nule.",
        "utjecaj missing ratea",
    )

    # ---------------- vizualne rekonstrukcije --------------------------------
    zamijeni_odlomak(
        doc,
        "Rekonstrukcijski prikazi služe za vizualnu provjeru onoga što se vidi u numeričkim metrikama. "
        "Originalni temperaturni niz uspoređen je s rekonstrukcijom, a maskirane pozicije označene su "
        "crvenim točkama. Prikazana je linearna interpolacija pri 20 % missing ratea za svih pet "
        "scenarija jer je ta metoda kroz cijeli eksperiment bila najstabilnija.",
        "Rekonstrukcijski prikazi služe za vizualnu provjeru onoga što se vidi u numeričkim metrikama. "
        "Originalni temperaturni niz uspoređen je s rekonstrukcijom, a maskirane pozicije označene su "
        "crvenim točkama. Prikazana je linearna interpolacija pri 20 % missing ratea za svih pet "
        "scenarija jer je ta metoda kroz cijeli eksperiment bila najstabilnija. Svi prikazi odnose se na "
        "prvi od 20 tjednih prozora, pa metrike navedene uz njih opisuju taj konkretni tjedan i "
        "očekivano odstupaju od prosjeka u tablicama.",
        "uvod u rekonstrukcije",
    )

    zamijeni_odlomak(
        doc,
        "U random missing scenariju uklonjene točke raspršene su po cijelom nizu, pa linearna "
        "interpolacija za većinu njih ima poznata mjerenja s obje strane. Na prikazu se zato "
        "rekonstruirana linija može usporediti s originalom upravo na pojedinačnim označenim pozicijama. "
        "Dobivene metrike potvrđuju vrlo dobru rekonstrukciju: MAE iznosi 0,0502, RMSE 0,0680, a R² "
        "0,9986.",
        "U random missing scenariju uklonjene točke raspršene su po cijelom nizu, pa linearna "
        "interpolacija za većinu njih ima poznata mjerenja s obje strane. Na prikazu se zato "
        "rekonstruirana linija može usporediti s originalom upravo na pojedinačnim označenim pozicijama. "
        f"Dobivene metrike potvrđuju vrlo dobru rekonstrukciju: MAE iznosi {R['random']['mae']}, RMSE "
        f"{R['random']['rmse']}, a R² {R['random']['r2']}.",
        "rekonstrukcija random",
    )

    zamijeni_odlomak(
        doc,
        "Pri 20 % missing ratea dobiveni su MAE 0,2849, RMSE 0,3440 i R² 0,8666. Rekonstrukcija i dalje "
        "prati glavni tijek niza, ali je odstupanje primjetno veće nego u random scenariju.",
        f"Pri 20 % missing ratea dobiveni su MAE {R['block']['mae']}, RMSE {R['block']['rmse']} i R² "
        f"{R['block']['r2']}. Rekonstrukcija i dalje prati glavni tijek niza, ali je odstupanje primjetno "
        "veće nego u random scenariju.",
        "rekonstrukcija block",
    )

    zamijeni_odlomak(
        doc,
        "Metrike potvrđuju da je taj slučaj problematičan: MAE je 1,1335, RMSE 1,3189, a R² -2,2624. "
        "Negativan R² pokazuje da rekonstrukcija na uklonjenom početnom dijelu slabo prati stvarne "
        "vrijednosti.",
        f"Metrike za taj slučaj iznose: MAE {R['block_start']['mae']}, RMSE {R['block_start']['rmse']}, a "
        f"R² {R['block_start']['r2']}. Pozitivan R² pokazuje da rekonstrukcija na uklonjenom početnom "
        "dijelu u ovom tjednu ipak prati stvarne vrijednosti, iako je pogreška veća nego u random "
        "scenariju.",
        "rekonstrukcija block_start",
    )

    zamijeni_odlomak(
        doc,
        "Kod block_middle scenarija praznina je smještena između poznatih dijelova niza, pa postoje rubne "
        "vrijednosti i prije i poslije bloka. Ipak, duljina kontinuirane praznine i dalje može dovesti do "
        "značajnog odstupanja unutar njezina središta. Za prikazani slučaj MAE iznosi 0,3654, RMSE 0,4325 "
        "i R² -2,2972. Iako su apsolutne pogreške manje nego kod block_start primjera, negativan R² "
        "pokazuje da rekonstruirane promjene na maskiranim točkama ne prate dovoljno dobro stvarnu "
        "varijabilnost.",
        "Kod block_middle scenarija praznina je smještena između poznatih dijelova niza, pa postoje rubne "
        "vrijednosti i prije i poslije bloka. Ipak, duljina kontinuirane praznine i dalje može dovesti do "
        f"značajnog odstupanja unutar njezina središta. Za prikazani slučaj MAE iznosi "
        f"{R['block_middle']['mae']}, RMSE {R['block_middle']['rmse']} i R² {R['block_middle']['r2']}. "
        "Negativan R² pokazuje da rekonstruirane promjene na maskiranim točkama ne prate dovoljno dobro "
        "stvarnu varijabilnost.",
        "rekonstrukcija block_middle",
    )

    zamijeni_odlomak(
        doc,
        "U block_end scenariju uklonjen je završni dio vremenskog niza, pa nakon praznine nema poznatih "
        "mjerenja koja bi ograničila rekonstrukciju s desne strane. Time se jasno vidi razlika između "
        "interpolacije unutar niza i procjene na njegovu rubu. Pri 20 % missing ratea linearna "
        "interpolacija ostvaruje MAE 0,3538, RMSE 0,3940 i R² 0,6565. Rezultat je slabiji od random "
        "scenarija, ali R² i dalje ostaje pozitivan, pa rekonstrukcija zadržava dio stvarnog obrasca "
        "promjene temperature.",
        "U block_end scenariju uklonjen je završni dio vremenskog niza, pa nakon praznine nema poznatih "
        "mjerenja koja bi ograničila rekonstrukciju s desne strane. Time se jasno vidi razlika između "
        "interpolacije unutar niza i procjene na njegovu rubu. Pri 20 % missing ratea linearna "
        f"interpolacija ostvaruje MAE {R['block_end']['mae']}, RMSE {R['block_end']['rmse']} i R² "
        f"{R['block_end']['r2']}. Od svih pet prikaza ovo je najslabiji rezultat, jer se bez desnog ruba "
        "posljednja poznata vrijednost produljuje kroz cijeli uklonjeni dio.",
        "rekonstrukcija block_end",
    )

    # ---------------- najbolja metoda po scenariju ---------------------------
    zamijeni_odlomak(
        doc,
        "Ako se kao kriterij uzme broj najboljih MAE rezultata, linear_interpolation uvjerljivo vodi s 27 "
        "pobjeda u 40 kombinacija. Cubic_interpolation je posebno jaka u random scenariju pri 10 %, 20 % "
        "i 30 %, dok se pri većim rateovima i u većini block scenarija češće pokazuje boljom linearna "
        "interpolacija. Time_interpolation i linear_interpolation daju jednake rezultate jer su mjerenja "
        "u korištenom Jena nizu raspoređena u pravilnim razmacima od 10 minuta. KNN ostaje prihvatljiv "
        "uglavnom kod random praznina, dok mu pogreška na kontinuiranim blokovima znatno raste. Decision "
        "Tree i Random Forest imaju nekoliko pojedinačnih pobjeda, ali kroz cijeli eksperimentalni skup "
        "ne pokazuju stabilnost linearne interpolacije.",
        "Ako se kao kriterij uzme broj najboljih MAE rezultata, linearna interpolacija vodi s 15 pobjeda u "
        "40 kombinacija, no te pobjede dijeli s vremenskom interpolacijom i KNN-om jer sve tri metode daju "
        "brojčano istovjetne vrijednosti. Time_interpolation i linear_interpolation podudaraju se zato što "
        "su mjerenja u korištenom Jena nizu raspoređena u pravilnim razmacima od 10 minuta, a KNN im se "
        "pridružuje jer u konačnoj izvedbi obuhvaća prazninu s obje strane i teži po inverznoj "
        "udaljenosti. Metode strojnog učenja zajedno imaju 20 pobjeda, sve odreda na kontinuiranim "
        "blokovima, pri čemu prednjači napredni KNN s osam. Kubna i spline interpolacija najbolje su samo "
        "pri najnižim rateovima u random scenariju te pri 10 % u block_start scenariju, a na duljim "
        "blokovima im pogreška naglo raste jer polinom trećeg stupnja izvan poznatih točaka divergira. "
        "Zbog izjednačenosti prve tri metode zbroj pobjeda po recima Tablice 5-7 veći je od 40.",
        "najbolja metoda po scenariju",
    )

    # ---------------- TUMACENJE REZULTATA ------------------------------------
    zamijeni_odlomak(
        doc,
        "Rezultati pokazuju da su klasične interpolacijske metode u ovom eksperimentu uspješnije od metoda "
        "strojnog učenja. Linearna i vremenska interpolacija imaju najniži prosječni MAE od 0,5855 i "
        "najmanju standardnu devijaciju MAE od 0,4353. To pokazuje da su stabilne kroz različite scenarije "
        "i missing rateove. Metode strojnog učenja nisu ostvarile ukupnu prednost, što se može objasniti "
        "kratkim vremenskim nizom od 288 zapisa i pravilnim 10-minutnim uzorkovanjem. U takvom nizu lokalna "
        "struktura i susjedne vrijednosti često daju dovoljno informacija za rekonstrukciju, pa jednostavne "
        "interpolacijske metode mogu nadmašiti složenije modele.",
        "Rezultati pokazuju da su linearna interpolacija i metode strojnog učenja u ovom eksperimentu vrlo "
        "blizu jedne drugima. Linearna i vremenska interpolacija imaju najniži prosječni MAE od 2,5121 i "
        "standardnu devijaciju MAE od 1,2775, a odmah iza njih slijede napredni KNN s 2,5232, neuronska "
        "mreža s 2,5354, slučajna šuma s 2,5440 i stablo odlučivanja s 2,5853. Razlika između prvog i "
        "petog mjesta iznosi manje od 0,08 °C, pa se nijedna od tih metoda ne može proglasiti uvjerljivo "
        "boljom. Jasno zaostaju samo pomični prosjek i forward fill, s 3,2492 i 3,3238, te kubna i spline "
        "interpolacija, koje s 8,5056 i 9,8159 na duljim blokovima potpuno gube stabilnost. Da metode "
        "strojnog učenja ne ostvaruju veću prednost može se objasniti pravilnim 10-minutnim uzorkovanjem, u "
        "kojem susjedne vrijednosti već sadrže gotovo svu informaciju potrebnu za rekonstrukciju.",
        "usporedba klasicnih i ML",
    )

    zamijeni_odlomak(
        doc,
        "Random missing scenarij najpovoljniji je za interpolacijske metode jer su uklonjene vrijednosti "
        "pojedinačno raspoređene kroz niz. Kubna interpolacija postiže najbolji MAE pri 10 %, 20 % i 30 %, "
        "s vrijednostima 0,0406, 0,0488 i 0,0448. Od 40 % do 80 % najbolja postaje linearna interpolacija, "
        "ali i tada pogreške ostaju vrlo male. Na random scenariju R² za najbolje metode ostaje iznad 0,996 "
        "i pri 80 % missing ratea. To znači da rekonstruirane vrijednosti i dalje dobro prate promjene "
        "originalnog temperaturnog niza, iako je većina vrijednosti umjetno uklonjena.",
        "Random missing scenarij najpovoljniji je za interpolacijske metode jer su uklonjene vrijednosti "
        "pojedinačno raspoređene kroz niz. Spline i kubna interpolacija postižu najbolji MAE pri 10 % do "
        "40 %, s vrijednostima od 0,0721 do 0,1055. Od 50 % do 80 % najbolja postaje linearna "
        "interpolacija, ali i tada pogreške ostaju vrlo male. Na random scenariju R² za najbolje metode "
        "ostaje iznad 0,99 i pri 80 % missing ratea. To znači da rekonstruirane vrijednosti i dalje dobro "
        "prate promjene originalnog temperaturnog niza, iako je većina vrijednosti umjetno uklonjena.",
        "tumacenje random",
    )

    zamijeni_odlomak(
        doc,
        "Block missing scenariji stvaraju teži problem jer unutar uklonjenog bloka nema poznatih "
        "vrijednosti. Linearna interpolacija je u takvim uvjetima često najbolja jer izravno povezuje "
        "poznate vrijednosti na rubovima bloka. Kod običnog block scenarija ona pobjeđuje u većini rateova, "
        "uključujući 80 % s MAE 0,6705. KNN je posebno slab kod block scenarija. Pri 50 % do 80 % missing "
        "ratea na block scenarijima MAE mu se kreće približno od 1,95 do 3,52 °C, a R² je često negativan. "
        "To pokazuje da vremenske značajke koje KNN koristi nisu dovoljne za pouzdanu rekonstrukciju dugih "
        "kontinuiranih praznina.",
        "Block missing scenariji stvaraju teži problem jer unutar uklonjenog bloka nema poznatih "
        "vrijednosti. Linearna interpolacija ondje daje solidan rezultat jer izravno povezuje poznate "
        "vrijednosti na rubovima bloka, ali je metode strojnog učenja u običnom block scenariju nadmašuju u "
        "svih osam rateova. Prednost je pritom sitna i na 80 % iznosi 0,0115 °C, što nije dovoljno da bi se "
        "govorilo o praktično značajnoj razlici. Znatno je važnije koje metode ondje potpuno zakažu: kubna "
        "i spline interpolacija na 80 % dosežu MAE od približno 29 °C jer polinom trećeg stupnja izvan "
        "poznatih točaka divergira, dok forward fill i pomični prosjek ostaju oko 5 °C jer unutar bloka "
        "nemaju nijednog poznatog susjeda i svode se na zadržavanje posljednje vrijednosti.",
        "tumacenje block",
    )

    zamijeni_odlomak(
        doc,
        "Pozicija bloka značajno utječe na rezultate. Block_start i block_end scenariji imaju problem jer se "
        "blok nalazi na rubu niza, pa metode nemaju poznate vrijednosti s obje strane praznine. Block_middle "
        "je povoljniji za interpolaciju pri nižim rateovima jer su poznate vrijednosti dostupne prije i "
        "poslije bloka. Pri 80 % missing ratea najteži scenarij je block_end s prosječnim MAE 3,4677, a odmah "
        "iza njega block_middle s 3,3376. Time se prošireni eksperiment razlikuje od ranije verzije s "
        "10–40 %, jer veliki rateovi jasnije pokazuju da rubni i veliki blokovi mogu značajno narušiti "
        "kvalitetu rekonstrukcije.",
        "Pozicija bloka značajno utječe na rezultate. Block_start i block_end scenariji imaju problem jer se "
        "blok nalazi na rubu niza, pa metode nemaju poznate vrijednosti s obje strane praznine. Block_middle "
        "je povoljniji za interpolaciju pri nižim rateovima jer su poznate vrijednosti dostupne prije i "
        f"poslije bloka. Pri 80 % missing ratea najveći prosječni MAE ima obični block scenarij, "
        f"{pr('block', '80', 'prosjek_mae')}, a odmah iza njega block_middle s "
        f"{pr('block_middle', '80', 'prosjek_mae')}. Ti prosjeci ipak više govore o metodama koje se "
        "raspadaju nego o težini samog scenarija, jer se najbolje metode u sva tri block scenarija na 80 % "
        "zadržavaju između 3,2 i 3,7 °C.",
        "tumacenje pozicija bloka",
    )

    zamijeni_odlomak(
        doc,
        "U završnom eksperimentu koristi se napredni KNN, registriran kao knn. Metoda koristi vremenske "
        "značajke, uključujući ciklički prikaz sata i dana u godini. Prednost takvog pristupa je mogućnost "
        "povezivanja vremenski sličnih uzoraka, primjerice sličnih sati ili dijelova godine. Glavno "
        "ograničenje KNN metode vidljivo je kod block scenarija i velikih missing rateova. KNN ne prati "
        "lokalni trend kroz dugi blok, pa pri 80 % missing ratea ostvaruje znatno veće pogreške od linearne "
        "interpolacije. Usporedba osnovnog i naprednog KNN-a nije moguća jer se u završnom CSV-u nalazi samo "
        "napredna verzija registrirana kao knn.",
        "U završnom eksperimentu obje su KNN izvedbe zasebno vrednovane. Osnovni KNN, registriran kao knn, "
        "traži poznate susjede s obje strane praznine i teži ih po inverznoj udaljenosti, zbog čega daje "
        "brojčano iste vrijednosti kao linearna interpolacija. Napredni KNN, registriran kao knn_upgraded, "
        "ne uspoređuje uzorke po položaju nego po obilježjima same praznine, primjerice po relativnom "
        "položaju unutar nje i udaljenosti do poznatih rubova, te uči odstupanje od linearne procjene. Takav "
        "pristup daje mu prednost upravo ondje gdje osnovni KNN nema što dodati: napredni KNN najbolja je "
        "metoda u osam od 40 kombinacija, sve na kontinuiranim blokovima, s prosječnim MAE 2,5232 naspram "
        "2,5121 osnovne izvedbe. Razlika je premala da bi se jedna izvedba mogla proglasiti boljom, ali "
        "pokazuje da dodatne značajke pomažu točno u onim uvjetima u kojima linearna procjena nema dovoljno "
        "oslonca.",
        "tumacenje KNN",
    )

    zamijeni_odlomak(
        doc,
        "Decision Tree i Random Forest mogu naučiti nelinearne odnose između vremenskih značajki i "
        "temperature. U pojedinim slučajevima ostvaruju najbolji rezultat, primjerice decision_tree u "
        "block_middle scenariju pri 10 % i random_forest u block_end scenariju pri 40 %. Ipak, te pobjede su "
        "iznimke, a ne pravilo. Prosječni MAE za Decision Tree iznosi 1,4954, dok Random Forest ima prosječni "
        "MAE 1,8691 u dijelu s većim missing rateovima od 50 % do 80 %. Obje metode su znatno manje stabilne "
        "od linearne interpolacije u ovom eksperimentalnom okviru.",
        "Decision Tree i Random Forest mogu naučiti nelinearne odnose između ulaznih značajki i temperature. "
        "U pojedinim slučajevima ostvaruju najbolji rezultat, primjerice stablo odlučivanja u block scenariju "
        "pri 70 % i 80 % te slučajna šuma u istom scenariju pri 40 % i 50 %. Prosječni MAE za stablo "
        "odlučivanja iznosi 2,5853, a za slučajnu šumu 2,5440, što ih smješta unutar 0,08 °C od linearne "
        "interpolacije. Obje su metode dakle konkurentne, ali ni jedna ne donosi prednost koja bi opravdala "
        "njihovu složenost na ovako pravilno uzorkovanom nizu. Njihova stvarna vrijednost pokazuje se tek na "
        "kontinuiranim blokovima, gdje linearna procjena nema dovoljno oslonca.",
        "tumacenje DT i RF",
    )

    zamijeni_odlomak(
        doc,
        "Glavno ograničenje rada je korištenje jednog temperaturnog niza iz Jena Climate skupa podataka. Niz "
        "sadrži 288 zapisa, odnosno 48 sati mjerenja u 10-minutnim intervalima, pa se zaključci prvenstveno "
        "odnose na ovaj konkretni vremenski raspon i tip podataka. Nedostajuće vrijednosti su umjetno "
        "generirane. Takav pristup omogućuje objektivnu evaluaciju jer su stvarne vrijednosti poznate, ali ne "
        "obuhvaća sve moguće uzroke i oblike stvarnih nedostajućih podataka. Također, moving average nije "
        "uključen u završni eksperimentalni CSV, a osnovni KNN nije uspoređen s naprednim KNN-om u istom "
        "skupu rezultata.",
        "Glavno ograničenje rada je korištenje jedne meteorološke varijable iz Jena Climate skupa podataka. "
        "Eksperiment se provodi nad 20 tjednih prozora od po 1008 zapisa, odnosno sedam dana mjerenja u "
        "10-minutnim intervalima, raspoređenih kroz godinu radi sezonske pokrivenosti. Zaključci se stoga "
        "odnose na temperaturu pravilno uzorkovanu u kratkim intervalima i ne moraju vrijediti za nizove s "
        "nepravilnim razmacima ili duljim razdobljima. Nedostajuće vrijednosti su umjetno generirane, što "
        "omogućuje objektivnu evaluaciju jer su stvarne vrijednosti poznate, ali ne obuhvaća sve moguće "
        "uzroke i oblike stvarnih nedostajućih podataka. Naposljetku, razlike među pet najboljih metoda manje "
        "su od 0,08 °C, pa poredak među njima treba uzimati s oprezom bez formalne provjere statističke "
        "značajnosti.",
        "ogranicenja rada",
    )

    zamijeni_odlomak(
        doc,
        "Makar u provedenim eksperimentima jednostavnije interpolacijske metode često postigle bolje "
        "numeričke rezultate, to ne znači da metode strojnog učenja nemaju važnu ulogu. Njihova prednost je u "
        "tome što mogu koristiti više ulaznih značajki, primjerice sat u danu, dan u godini i položaj mjerenja "
        "u nizu, te na temelju njih učiti složenije obrasce u podacima. Takav pristup može biti korisniji kod "
        "većih i raznolikijih skupova podataka, nepravilnih vremenskih razmaka ili situacija u kojima "
        "vrijednost ne ovisi samo o najbližim poznatim mjerenjima, nego i o širem kontekstu vremenskog niza.",
        "Iako jednostavna linearna interpolacija ostaje ukupno najbolja, razlika prema metodama strojnog "
        "učenja svela se na nekoliko stotinki stupnja, a na kontinuiranim blokovima prednost prelazi na "
        "stranu strojnog učenja. Njihova je prednost u tome što mogu koristiti više ulaznih značajki, "
        "primjerice sat u danu, dan u godini, položaj mjerenja u nizu te udaljenost do poznatih vrijednosti s "
        "obje strane praznine, i na temelju njih učiti odstupanje od jednostavne linearne procjene. Takav "
        "pristup može biti korisniji kod većih i raznolikijih skupova podataka, nepravilnih vremenskih "
        "razmaka ili situacija u kojima vrijednost ne ovisi samo o najbližim poznatim mjerenjima, nego i o "
        "širem kontekstu vremenskog niza.",
        "uloga ML",
    )

    # ---------------- ZAKLJUCAK ----------------------------------------------
    zamijeni_odlomak(
        doc,
        "Konačni doprinos rada nije samo implementacija pojedinih metoda, nego i usporedbeni eksperimentalni "
        "okvir u kojem se iste metode testiraju nad istim oštećenim nizovima i istim maskama. Provedeno je "
        "ukupno 320 eksperimenata, što uključuje osam metoda, pet scenarija i missing rateove od 10 % do "
        "80 %. Rezultati su vrednovani metrikama MAE, RMSE i R² samo na pozicijama koje su umjetno uklonjene.",
        "Konačni doprinos rada nije samo implementacija pojedinih metoda, nego i usporedbeni eksperimentalni "
        "okvir u kojem se iste metode testiraju nad istim oštećenim nizovima i istim maskama. Provedeno je "
        f"ukupno {b['n_kombinacija'] * b['n_metoda']} kombinacija, što uključuje {n_met} metoda, pet "
        f"scenarija i missing rateove od 10 % do 80 %, a svaka je kombinacija ponovljena nad "
        f"{b['n_ponavljanja']} neovisnih tjednih prozora, ukupno {b['n_izvodenja']} izvođenja. Rezultati su "
        "vrednovani metrikama MAE, RMSE i R² samo na pozicijama koje su umjetno uklonjene.",
        "zakljucak doprinos",
    )

    zamijeni_odlomak(
        doc,
        "Rezultati pokazuju da jednostavnije metode nisu nužno lošije od složenijih metoda strojnog učenja. "
        "Linearna interpolacija pokazala se kao najstabilniji ukupni pristup jer ostvaruje najbolji MAE u "
        "najvećem broju kombinacija i ima najniži prosječni MAE. Vremenska interpolacija daje iste rezultate "
        "kao linearna interpolacija jer su mjerenja u korištenom nizu ravnomjerno raspoređena u 10-minutnim "
        "intervalima.",
        "Rezultati pokazuju da jednostavnije metode nisu nužno lošije od složenijih metoda strojnog učenja. "
        "Linearna interpolacija pokazala se kao najstabilniji ukupni pristup jer ostvaruje najbolji MAE u "
        "najvećem broju kombinacija i ima najniži prosječni MAE, ali je prednost pred najboljim metodama "
        "strojnog učenja manja od 0,08 °C i na kontinuiranim blokovima nestaje. Vremenska interpolacija daje "
        "iste rezultate kao linearna interpolacija jer su mjerenja u korištenom nizu ravnomjerno raspoređena "
        "u 10-minutnim intervalima, a isto vrijedi i za osnovni KNN jer obuhvaća prazninu s obje strane.",
        "zakljucak rezultati",
    )

    zamijeni_odlomak(
        doc,
        "Kubna interpolacija pokazala se vrlo dobrom kod random missing scenarija pri nižim missing "
        "rateovima, dok se kod većih i rubnih blokova prednost češće prebacuje na linearnu interpolaciju. "
        "Metode strojnog učenja, uključujući KNN, Decision Tree i Random Forest, povremeno ostvaruju dobre "
        "rezultate u pojedinim scenarijima, ali nisu pokazale jednaku stabilnost kroz cijeli eksperimentalni "
        "skup.",
        "Kubna i spline interpolacija pokazale su se vrlo dobrima kod random missing scenarija pri nižim "
        "missing rateovima, ali na kontinuiranim blokovima daju daleko najveće pogreške od svih promatranih "
        "metoda. Metode strojnog učenja, uključujući napredni KNN, stablo odlučivanja, slučajnu šumu i "
        "neuronsku mrežu, najbolje su u 20 od 40 kombinacija, gotovo isključivo na kontinuiranim blokovima, "
        "i kroz cijeli su skup jednako stabilne kao linearna interpolacija.",
        "zakljucak kubna i ML",
    )

    zamijeni_odlomak(
        doc,
        "Budući rad mogao bi uključiti dulje vremenske nizove, više meteoroloških varijabli, stvarne "
        "nedostajuće podatke, implementaciju pomičnog prosjeka te usporedbu osnovnog i naprednog KNN-a u "
        "istom eksperimentalnom okviru.",
        "Budući rad mogao bi uključiti dulje vremenske nizove, više meteoroloških varijabli, stvarne "
        "nedostajuće podatke te formalnu provjeru statističke značajnosti razlika među vodećim metodama, "
        "budući da su one u ovom eksperimentu manje od 0,08 °C.",
        "zakljucak buduci rad",
    )

    doc.save(str(IZLAZ))

    print(f"\nzamjena odlomaka: {odradeno} uspjesnih")
    if promasaji:
        print(f"NEPRONADENO ({len(promasaji)}):")
        for p in promasaji:
            print(f"   - {p}")
    else:
        print("svi ciljani odlomci pronadeni i zamijenjeni")
    print(f"\nspremljeno: {IZLAZ.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
