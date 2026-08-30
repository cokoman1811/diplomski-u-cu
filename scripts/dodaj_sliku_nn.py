#!/usr/bin/env python3
"""
Umece dijagram neuronske mreze u teorijski odjeljak, s natpisom u istom
obliku kao ostale metode (stil Caption, polja STYLEREF i SEQ).

Idempotentno: ako natpis vec postoji, nista ne dira.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, str(Path(__file__).resolve().parent))
from azuriraj_word import norm  # noqa: E402
from generiraj_dijagram_nn import IZLAZ as SLIKA, main as nacrta  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
DOK = ROOT / "docs" / "Diplomski-Toni_Jakelic_azurirano.docx"
NATPIS = ("Dijagram rada neuronske mreže. Izvor: vlastita izrada "
          "prema implementaciji projekta.")
# Sirina kao dijagram Decision Tree (wp:extent cx = 4635847 EMU).
SIRINA = Emu(4635847)


def nadi_pocetak(doc, prefiks: str):
    cilj = norm(prefiks)
    for par in doc.paragraphs:
        if norm(par.text).startswith(cilj):
            return par
    return None


def ocisti_oznake(el) -> None:
    for tag in ("w:bookmarkStart", "w:bookmarkEnd", "w:hyperlink",
                "w:proofErr", "w:commentRangeStart", "w:commentRangeEnd"):
        for e in list(el.iter(qn(tag))):
            roditelj = e.getparent()
            if roditelj is None:
                continue
            if tag == "w:hyperlink":
                # Djeca hyperlinka ostaju u odlomku, sam omotac nestaje.
                idx = list(roditelj).index(e)
                for i, child in enumerate(list(e)):
                    e.remove(child)
                    roditelj.insert(idx + i, child)
            roditelj.remove(e)


def umetni_sliku(doc) -> bool:
    if any(NATPIS in p.text for p in doc.paragraphs):
        print("dijagram vec postoji, preskacem")
        return False

    sidro = nadi_pocetak(doc, "Neuronska mreža procjenjuje temperaturu")
    uzor_natpis = nadi_pocetak(doc, "Slika")
    # Uzmi natpis Decision Tree, da SEQ ostane u istom poglavlju.
    for p in doc.paragraphs:
        if "Dijagram rada Decision Tree metode" in p.text:
            uzor_natpis = p
            break
    if sidro is None or uzor_natpis is None:
        raise RuntimeError("Ne mogu pronaci sidro ili uzorak natpisa.")
    if not SLIKA.exists():
        nacrta()

    # Odlomak za sliku: klon Normal odlomka, pa slika umjesto teksta.
    slika_el = copy.deepcopy(sidro._element)
    ocisti_oznake(slika_el)
    for r in list(slika_el.iter(qn("w:r"))):
        roditelj = r.getparent()
        if roditelj is not None:
            roditelj.remove(r)
    ppr = slika_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = slika_el.makeelement(qn("w:pPr"), {})
        slika_el.insert(0, ppr)
    for child in list(ppr):
        if child.tag in (qn("w:jc"), qn("w:spacing")):
            ppr.remove(child)
    spacing = ppr.makeelement(qn("w:spacing"),
                              {qn("w:before"): "120", qn("w:after"): "60"})
    jc = ppr.makeelement(qn("w:jc"), {qn("w:val"): "center"})
    ppr.append(spacing)
    ppr.append(jc)
    sidro._element.addnext(slika_el)
    slika_par = Paragraph(slika_el, sidro._parent)
    slika_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = slika_par.add_run()
    run.add_picture(str(SLIKA), width=SIRINA)

    # Natpis: klon s poljima STYLEREF i SEQ, bez knjiznih oznaka.
    cap_el = copy.deepcopy(uzor_natpis._element)
    ocisti_oznake(cap_el)
    slika_el.addnext(cap_el)
    cap_par = Paragraph(cap_el, uzor_natpis._parent)
    zamijenjeno = False
    for t in cap_par._element.iter(qn("w:t")):
        if t.text and "Dijagram rada Decision Tree metode" in t.text:
            t.text = t.text.replace(
                "Dijagram rada Decision Tree metode",
                "Dijagram rada neuronske mreže",
            )
            zamijenjeno = True
    if not zamijenjeno:
        # Fallback: upisi cijeli natpis u zadnji w:t.
        cvorovi = list(cap_par._element.iter(qn("w:t")))
        if cvorovi:
            cvorovi[-1].text = " " + NATPIS
    return True


def main() -> None:
    if not DOK.exists():
        sys.exit(f"Nema {DOK}. Prvo pokreni azuriraj_word.py i dopuni_poglavlja.py")
    if not SLIKA.exists():
        nacrta()
    doc = docx.Document(str(DOK))
    if umetni_sliku(doc):
        doc.save(str(DOK))
        print(f"slika umetnuta, spremljeno: {DOK.relative_to(ROOT)}")
    else:
        print("nema izmjena")


if __name__ == "__main__":
    main()
