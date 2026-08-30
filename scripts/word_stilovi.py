#!/usr/bin/env python3
"""Ispisuje stilove koje dokument stvarno koristi, radi vjernog opona\u0161anja."""

from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
PUT = ROOT / "docs" / "Diplomski-Toni_Jakelic_20.8_1.docx"


def blokovi(parent):
    for child in parent.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def main() -> None:
    d = docx.Document(str(PUT))

    print("=" * 88)
    print("UPOTREBA STILOVA U ODLOMCIMA")
    print("=" * 88)
    c = Counter(p.style.name for p in d.paragraphs)
    for ime, n in c.most_common():
        print(f"  {ime:<34}{n:>5}")

    print()
    print("=" * 88)
    print("STILOVI TABLICA")
    print("=" * 88)
    for i, t in enumerate(d.tables, 1):
        print(f"  [{i}] {t.style.name if t.style else '(bez stila)'}")

    print()
    print("=" * 88)
    print("NATPISI: odlomci koji po\u010dinju s 'Tablica' ili 'Slika'")
    print("=" * 88)
    prosli = None
    for b in blokovi(d):
        if isinstance(b, Table):
            prosli = "TABLICA"
            continue
        t = b.text.strip()
        if t.startswith(("Tablica", "Slika")):
            polozaj = "odmah nakon tablice" if prosli == "TABLICA" else ""
            print(f"  [{b.style.name}] {t[:88]}   {polozaj}")
        if t:
            prosli = "P"

    print()
    print("=" * 88)
    print("FORMAT BROJEVA: uzorci iz teksta")
    print("=" * 88)
    import re
    uzorci = Counter()
    for p in d.paragraphs:
        for m in re.finditer(r"\d+[.,]\d+", p.text):
            s = m.group()
            uzorci["zarez" if "," in s else "tocka"] += 1
            uzorci[f"decimala_{len(s.split(',' if ',' in s else '.')[1])}"] += 1
    for k, v in sorted(uzorci.items()):
        print(f"  {k:<20}{v:>5}")

    print()
    print("=" * 88)
    print("PRVA TABLICA \u2014 formatiranje \u0107elija zaglavlja i prvog retka")
    print("=" * 88)
    t = d.tables[2]
    for r in range(min(2, len(t.rows))):
        for cl in t.rows[r].cells[:3]:
            for par in cl.paragraphs:
                bold = [run.bold for run in par.runs]
                font = [(run.font.name, run.font.size) for run in par.runs]
                print(f"  red {r} | '{par.text[:22]:<22}' stil={par.style.name} "
                      f"bold={bold} font={font}")


if __name__ == "__main__":
    main()
